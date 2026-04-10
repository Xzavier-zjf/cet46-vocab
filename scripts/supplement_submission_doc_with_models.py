from __future__ import annotations

from pathlib import Path
import shutil
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")
PDF_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.pdf")
BACKUP_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版-补充前备份.docx")
DOWNLOAD_DIAGRAM_DIR = Path(r"D:/zjf20/Downloads")
REPO_DIAGRAM_DIR = Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/diagrams")


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_run_fonts(run, size_pt=10.5, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def format_paragraph(paragraph, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, size_pt=size_pt, bold=run.bold if run.bold else False)


def add_text_before(anchor, text, style=None, size_pt=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    p = anchor.insert_paragraph_before("", style=style)
    run = p.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(p, size_pt=size_pt, align=align, first_line_indent=first_line_indent)
    return p


def add_caption_before(anchor, text, placeholder=False):
    return add_text_before(anchor, text, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)


def move_table_before(anchor, table):
    tbl = table._tbl
    anchor._p.addprevious(tbl)
    return table


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(p, size_pt=size_pt, align=align, first_line_indent=False)


def build_table(doc, rows, cols, widths_cm=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    if widths_cm:
        for idx, width in enumerate(widths_cm):
            for cell in table.columns[idx].cells:
                cell.width = Cm(width)
    return table


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_paragraph_index(doc, text):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise RuntimeError(f"Paragraph not found: {text}")


def remove_inserted_range(doc, start_text, end_text):
    try:
        start_idx = find_paragraph_index(doc, start_text)
        end_idx = find_paragraph_index(doc, end_text)
    except RuntimeError:
        return
    for idx in range(end_idx - 1, start_idx - 1, -1):
        delete_paragraph(doc.paragraphs[idx])


def remove_range_between(doc, start_heading, end_heading, keep_first_n=0):
    start_idx = find_paragraph_index(doc, start_heading)
    end_idx = find_paragraph_index(doc, end_heading)
    for idx in range(end_idx - 1, start_idx + keep_first_n, -1):
        delete_paragraph(doc.paragraphs[idx])


def add_table_title_and_table_before(doc, anchor, title, headers, rows, widths_cm):
    add_caption_before(anchor, title)
    table = build_table(doc, rows=len(rows) + 1, cols=len(headers), widths_cm=widths_cm)
    move_table_before(anchor, table)
    for col_idx, header in enumerate(headers):
        set_cell_text(table.cell(0, col_idx), header, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(row_idx, col_idx), value, bold=False, align=align)
    return table


def mx_geometry(x, y, w, h, relative=False, points=None):
    geom = ET.Element("mxGeometry", {"as": "geometry"})
    if relative:
        geom.set("relative", "1")
    else:
        geom.set("x", str(x))
        geom.set("y", str(y))
        geom.set("width", str(w))
        geom.set("height", str(h))
    if points:
        array = ET.SubElement(geom, "Array", {"as": "points"})
        for px, py in points:
            ET.SubElement(array, "mxPoint", {"x": str(px), "y": str(py)})
    return geom


def cell(cell_id, value="", style="", vertex=False, edge=False, parent="1", source=None, target=None, geometry=None):
    attrs = {"id": str(cell_id), "value": value, "style": style, "parent": parent}
    if vertex:
        attrs["vertex"] = "1"
    if edge:
        attrs["edge"] = "1"
    if source:
        attrs["source"] = str(source)
    if target:
        attrs["target"] = str(target)
    c = ET.Element("mxCell", attrs)
    c.append(geometry if geometry is not None else mx_geometry(0, 0, 0, 0))
    return c


def drawio_file(name, cells):
    mxfile = ET.Element("mxfile", {"host": "app.diagrams.net", "modified": "2026-04-06T00:00:00.000Z", "agent": "Codex", "version": "24.7.17"})
    diagram = ET.SubElement(mxfile, "diagram", {"id": name, "name": "Page-1"})
    model = ET.SubElement(diagram, "mxGraphModel", {
        "dx": "1650",
        "dy": "900",
        "grid": "1",
        "gridSize": "10",
        "guides": "1",
        "tooltips": "1",
        "connect": "1",
        "arrows": "1",
        "fold": "1",
        "page": "1",
        "pageScale": "1",
        "pageWidth": "1654",
        "pageHeight": "1169",
        "math": "0",
        "shadow": "0",
    })
    root = ET.SubElement(model, "root")
    root.append(ET.Element("mxCell", {"id": "0"}))
    root.append(ET.Element("mxCell", {"id": "1", "parent": "0"}))
    for c in cells:
        root.append(c)
    return ET.tostring(mxfile, encoding="unicode")


def write_text(path, content):
    path.write_text(content, encoding="utf-8")


def point(x, y, as_name=None):
    attrs = {"x": str(x), "y": str(y)}
    if as_name:
        attrs["as"] = as_name
    return ET.Element("mxPoint", attrs)


def point_edge(cell_id, value, style, sx, sy, tx, ty, points=None, parent="1"):
    edge = ET.Element("mxCell", {
        "id": str(cell_id),
        "value": value,
        "style": style,
        "parent": parent,
        "edge": "1",
    })
    geom = ET.Element("mxGeometry", {"relative": "1", "as": "geometry"})
    geom.append(point(sx, sy, "sourcePoint"))
    geom.append(point(tx, ty, "targetPoint"))
    if points:
        array = ET.SubElement(geom, "Array", {"as": "points"})
        for px, py in points:
            array.append(point(px, py))
    edge.append(geom)
    return edge


def insert_section_34(doc):
    remove_range_between(doc, "3.4 关键业务流程分析", "3.5 本章小结", keep_first_n=1)
    anchor = find_paragraph(doc, "3.5 本章小结")

    add_text_before(anchor, "结合当前项目的页面交互、接口调用、缓存更新与模型生成链路，系统关键业务流程不再仅以概述方式展开，而是进一步从用例关系、数据流向与活动步骤三个层面进行细化。这样既能够增强需求分析的可读性，也能为后续总体设计与详细实现提供更稳定的建模依据。")

    add_text_before(anchor, "3.4.1 系统用例图", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor, "系统主要参与者与核心功能边界如图3-1所示。该用例图围绕普通用户、管理员、本地模型服务和云端模型服务四类主体展开，重点反映学习端与管理端围绕词汇学习闭环所形成的交互关系。")
    add_caption_before(anchor, "（此处插入图3-1 系统用例图）", placeholder=True)
    add_caption_before(anchor, "图3-1 系统用例图")
    add_text_before(anchor, "由图3-1可见，普通用户主要围绕注册登录、词库浏览、详情学习、加入学习计划、AI 内容生成、复习、测验、学习助手和看板查看等能力展开；管理员则额外承担词库导入回滚、模型管理等后台职责。本地模型服务与云端模型服务不直接面向最终用户，而是以支撑角色参与解释生成和对话生成流程，从而保证系统在智能生成链路上的可扩展性与连续性。")

    add_text_before(anchor, "3.4.2 系统关键用例表", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor, "为了进一步说明关键业务场景的触发条件、执行路径和异常处理方式，本文将系统中最具代表性的用户侧与管理侧用例归纳为表3-1。该表兼顾正常流程和备选流程，有助于后续系统设计与测试章节形成对应关系。")
    use_case_rows = [
        ("UC-01", "用户注册与登录", "普通用户", "用户未登录，系统服务可用", "填写注册信息并完成登录，请求 /auth/register 与 /auth/login 建立账号和会话", "用户名重复或密码不合规时返回错误提示，登录失败时拒绝签发令牌", "用户获得 JWT 并进入学习首页"),
        ("UC-02", "浏览词库并加入学习", "普通用户", "用户已登录，词库数据已导入", "请求 /word/list 浏览词条，再通过 /word/learn/add 将目标单词加入学习计划", "词条不存在或已加入计划时给出提示", "生成 user_word_progress 初始记录"),
        ("UC-03", "查看单词详情并触发 AI 生成", "普通用户", "用户已登录，词条存在", "先请求 /word/detail 加载基础信息，再调用 /word/llm/generate 或解释生成接口补齐扩展内容", "缓存未命中、模型超时或返回不完整时进入 pending、partial 或 fallback 分支", "word_meta 与缓存内容被刷新，页面展示最终结果"),
        ("UC-04", "完成每日复习", "普通用户", "存在 next_review_date 到期词汇", "请求 /review/today 获取卡片，完成评分后调用 /review/submit 更新调度参数", "无待复习单词时返回空列表，非法评分时拒绝提交", "学习进度、复习日志与看板缓存同步更新"),
        ("UC-05", "发起并提交模拟测验", "普通用户", "用户已登录且词表中存在足够词汇", "调用 /quiz/generate 生成题目，答题后调用 /quiz/submit 保存结果", "quizId 无效或答案缺失时拒绝提交", "生成 quiz_session_record 历史记录并返回得分"),
        ("UC-06", "使用学习助手问答", "普通用户", "用户已登录，助手会话状态可读", "通过 /assistant/state 读取会话上下文，再调用 /assistant/chat 发起问答", "模型输出中断时执行续写或回退，状态同步失败时保留已落库内容", "会话消息与上下文被持久化，页面展示回复"),
        ("UC-07", "管理员导入词库", "管理员", "管理员已登录，上传词库文件格式正确", "先用 /admin/word-bank/preview 预览导入结果，再通过 /admin/word-bank/import 执行正式导入", "字段缺失、格式错误或导入冲突时标记为 skipped 或返回失败原因", "写入词表、导入批次和批次明细"),
        ("UC-08", "管理员配置云模型", "管理员", "管理员已登录，具备模型维护权限", "调用 /admin/llm/cloud-models 新增、修改或设置默认模型", "模型键冲突、协议不合法或 API Key 无效时返回失败", "cloud_llm_model 配置被更新并影响后续生成链路"),
    ]
    add_table_title_and_table_before(doc, anchor, "表3-1 系统关键用例说明表", ["用例编号", "用例名称", "参与者", "前置条件", "基本流程", "异常/备选流程", "后置结果"], use_case_rows, [1.8, 2.6, 1.8, 3.2, 4.6, 4.2, 3.2])
    add_text_before(anchor, "从表3-1可以看出，系统关键用例既覆盖了普通用户的学习闭环，也覆盖了管理员围绕词库与模型配置进行的后台维护行为。各个用例在前置条件、异常路径和后置结果上均与实际接口和数据结构保持一致，因此可直接为后文的接口设计、数据库设计和测试用例设计提供依据。")

    add_text_before(anchor, "3.4.3 数据流图", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor, "系统输入、处理、存储与输出之间的主要关系如图3-2所示。该数据流图以用户和管理员为输入源，以本地模型和云端模型为外部智能服务，以认证授权、词库学习、AI 内容生成、复习调度、测验管理、学习助手和后台词库管理为处理中心，对核心数据在系统内的流向进行了整合表达。")
    add_caption_before(anchor, "（此处插入图3-2 数据流图）", placeholder=True)
    add_caption_before(anchor, "图3-2 数据流图")
    add_text_before(anchor, "由图3-2可见，普通业务链路主要围绕用户请求、业务处理、数据库读写和结果返回展开；AI 生成链路则在此基础上额外引入模型调用、结构化解析、结果回填和缓存刷新等步骤。也就是说，系统并非把大模型当作独立模块孤立调用，而是将其嵌入到词条详情、学习助手等既有业务流程之中，从而形成可缓存、可追踪、可重试的工程化数据流。")

    add_text_before(anchor, "3.4.4 活动图", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor, "用户查看单词详情并触发 AI 解释生成的活动过程如图3-3所示。该图聚焦于系统最具特色的智能生成链路，强调缓存命中、数据库读取、异步生成、结果解析与失败降级之间的先后关系。")
    add_caption_before(anchor, "（此处插入图3-3 活动图）", placeholder=True)
    add_caption_before(anchor, "图3-3 活动图")
    add_text_before(anchor, "由图3-3可见，系统在用户进入详情页后并不会直接无条件调用模型，而是先检查缓存和已有元数据状态；只有在缓存未命中、数据缺失或状态未完成时，才进入模型生成和回填流程。这样的活动路径既降低了重复生成成本，也使系统在异常情况下能够返回基础词义或降级结果，从而保证学习页面的可用性。")


def insert_section_43_and_44(doc):
    remove_inserted_range(doc, "4.3.1 主要数据表设计", "4.4 接口设计")
    remove_inserted_range(doc, "4.4.1 系统核心接口设计表", "4.5 安全、缓存与异常处理设计")

    anchor_44 = find_paragraph(doc, "4.4 接口设计")
    add_text_before(anchor_44, "4.3.1 主要数据表设计", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor_44, "从总体结构看，系统数据库围绕用户、词汇、学习进度、行为日志、测验结果、模型配置以及导入回滚记录展开。其中，基础词表承担静态词汇资源存储职责，扩展元数据与日志表承担动态学习数据和 AI 生成数据的存储职责，后台管理表则用于支撑模型治理和词库维护。")
    overview_rows = [
        ("user", "保存账号、角色、学习风格和模型偏好", "username、role、llm_style、llm_provider", "与 user_word_progress、review_log、quiz_session_record 关联"),
        ("cet4zx / cet6zx", "保存基础词汇资源", "id、english、sent、chinese", "作为 word_meta、user_word_progress、review_log 的词汇来源"),
        ("word_meta", "保存 AI 生成扩展内容", "word_id、word_type、style、ai_explain、gen_status", "与基础词表按词汇编号和词库类型关联"),
        ("user_word_progress", "保存复习调度状态", "user_id、word_id、word_type、easiness、next_review_date", "与 user、词表、review_log 协同工作"),
        ("review_log", "保存复习行为日志", "user_id、word_id、word_type、score、reviewed_at", "支撑看板统计与复习追踪"),
        ("quiz_session_record", "保存测验历史结果", "quiz_id、user_id、quiz_type、total、correct、details_json", "与 user 构成一对多关系"),
        ("cloud_llm_model", "保存云模型接入配置", "provider、model_key、base_url、protocol、visibility", "支撑学习助手与 AI 解释生成"),
        ("word_import_batch / item", "保存词库导入与回滚明细", "batch_id、word_type、action_type、rolled_back", "支撑后台导入审计与回滚"),
    ]
    add_table_title_and_table_before(doc, anchor_44, "表4-1 主要业务数据表概览", ["表名", "作用", "关键字段", "与其他表关系"], overview_rows, [3.0, 4.2, 5.0, 5.0])

    core_table_rows = [
        ("user", "id；username；password；role；llm_style；llm_provider；llm_local_model；llm_cloud_model；daily_target", "主键 id，用户名用于账号识别，角色与模型偏好共同影响权限和生成策略", "承载账号认证、风格选择、每日目标与模型偏好"),
        ("word_meta", "id；word_id；word_type；style；sentence_en；sentence_zh；mnemonic；root_analysis；ai_explain；gen_status", "同一词条可按不同 style 生成多组扩展内容，gen_status 与 ai_explain_status 用于跟踪异步状态", "承载单词详情页和解释生成链路的核心元数据"),
        ("user_word_progress", "id；user_id；word_id；word_type；easiness；interval；repetition；next_review_date", "按用户、词库类型和词汇编号联合定位一条学习进度记录", "承载 SM-2 间隔重复调度参数"),
        ("review_log", "id；user_id；word_id；word_type；score；time_spent_ms；reviewed_at", "每次复习行为单独记录，便于统计趋势与行为回溯", "承载复习日志、看板统计与行为分析"),
        ("quiz_session_record", "id；quiz_id；user_id；word_type；quiz_type；total；correct；wrong_count；details_json；created_at", "以 quiz_id 和详情 JSON 保存测验快照，避免题目内容丢失", "承载测验得分、错题回看和历史详情展示"),
    ]
    add_table_title_and_table_before(doc, anchor_44, "表4-2 核心学习业务数据表设计", ["表名", "主要字段", "关键约束/状态", "说明"], core_table_rows, [2.6, 6.3, 4.5, 4.2])

    admin_table_rows = [
        ("cloud_llm_model", "id；provider；model_key；base_url；path；protocol；display_name；enabled；is_default；visibility；owner_user_id", "通过 visibility 与 owner_user_id 实现全局模型与私有模型隔离", "保存可被系统调用的云模型配置"),
        ("cloud_llm_provider_credential", "id；provider；visibility；owner_user_id；api_key_ciphertext；api_key_mask", "同一作用域下同一 provider 仅允许保留一份凭据", "保存云模型提供商级别的访问凭据"),
        ("word_import_batch", "batch_id；word_type；file_name；inserted_count；updated_count；skipped_count；status；created_by", "以 batch_id 作为批次主键，记录导入状态与统计结果", "保存一次导入任务的总体信息"),
        ("word_import_batch_item", "id；batch_id；word_type；word_id；action_type；old_*；new_*；rolled_back；created_at", "通过 batch_id 与批次表关联，并记录回滚标记", "保存导入过程中的逐词条明细"),
    ]
    add_table_title_and_table_before(doc, anchor_44, "表4-3 后台管理与模型配置数据表设计", ["表名", "主要字段", "关键约束/状态", "说明"], admin_table_rows, [3.1, 6.0, 4.6, 4.0])
    add_text_before(anchor_44, "上述设计说明系统数据库并不是简单堆叠若干业务表，而是通过基础词表、扩展元数据、学习行为记录和后台管理配置的分层协作来支撑完整功能。对于学习链路而言，`word_meta` 与 `user_word_progress` 是核心；对于后台治理而言，`cloud_llm_model` 与导入批次表则保证了配置可控与过程可审计。")

    add_text_before(anchor_44, "4.3.2 关键约束与索引设计", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor_44, "为了兼顾查询效率、作用域隔离和后台审计能力，系统在学习进度、模型配置和导入日志等关键对象上设置了有针对性的约束与索引，如表4-4所示。")
    index_rows = [
        ("user_word_progress", "idx_user_word_progress_user_word_type", "联合索引", "(user_id, word_type, word_id)", "快速定位某个用户在指定词库中的学习进度，避免重复插入"),
        ("cloud_llm_model", "uk_cloud_llm_model_visibility_owner_model", "唯一约束", "(visibility, owner_user_id, model_key)", "保证全局模型与私有模型在各自作用域内不发生键冲突"),
        ("cloud_llm_provider_credential", "uk_cloud_llm_provider_credential_scope_provider", "唯一约束", "(visibility, owner_user_id, provider)", "保证同一作用域下同一提供商只有一份有效凭据"),
        ("rbac_role_permission", "PRIMARY KEY", "复合主键", "(role_code, permission_code)", "防止同一角色重复绑定同一权限"),
        ("word_import_batch_item", "idx_word_import_batch_item_batch", "普通索引", "(batch_id)", "加快按批次查询导入明细与回滚项的速度"),
        ("rbac_role_permission_audit", "idx_rbac_role_permission_audit_changed_at", "普通索引", "(changed_at)", "支持按时间维度检索权限变更审计记录"),
    ]
    add_table_title_and_table_before(doc, anchor_44, "表4-4 关键约束与索引设计表", ["对象", "约束/索引名称", "类型", "字段组合", "设计目的"], index_rows, [3.0, 4.9, 2.4, 4.6, 4.1])
    add_text_before(anchor_44, "从表4-4可以看出，学习进度表采用联合索引是因为复习卡片、详情状态和加入学习操作都会高频按“用户 + 词库类型 + 词汇编号”定位记录；云模型配置采用作用域唯一约束，则是为了兼顾全局模型共享与私有模型隔离；导入批次和权限审计相关索引则主要服务于后台查询、回滚和审计追踪场景。")

    anchor_45 = find_paragraph(doc, "4.5 安全、缓存与异常处理设计")
    add_text_before(anchor_45, "4.4.1 系统核心接口设计表", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor_45, "在统一 REST 风格的基础上，系统将用户端学习接口、测评接口、学习助手接口与管理员维护接口分层组织。为突出论文中最具代表性的接口，表4-5对核心接口的路径、权限、请求参数和返回内容进行了集中整理。")
    interface_rows = [
        ("认证", "/auth/register", "POST", "匿名", "username、password", "用户基本信息与登录态初始化结果", "注册账号并建立用户基础信息"),
        ("认证", "/auth/login", "POST", "匿名", "username、password", "JWT、用户角色、基础资料", "完成身份认证并返回登录令牌"),
        ("词汇学习", "/word/list", "GET", "普通用户", "wordType、keyword、page、size", "分页词条列表", "按词库和关键词浏览词汇"),
        ("词汇学习", "/word/detail", "GET", "普通用户", "wordId、wordType、style", "基础释义与扩展元数据", "查看单词详情及生成状态"),
        ("AI 生成", "/word/llm/generate", "POST", "普通用户", "wordId、wordType、style、promptType", "任务触发结果与状态标记", "触发例句、近义词、助记等生成"),
        ("复习", "/review/today", "GET", "普通用户", "wordType、limit", "今日待复习词条列表", "读取当日复习任务"),
        ("复习", "/review/submit", "POST", "普通用户", "wordId、wordType、score、timeSpentMs", "更新后的进度与统计信息", "提交复习评分并更新 SM-2 调度"),
        ("测验", "/quiz/generate", "POST", "普通用户", "wordType、mode、count", "quizId 与题目集合", "生成选择题或填空题测验"),
        ("测验", "/quiz/submit", "POST", "普通用户", "quizId、answers", "总题数、正确数、错题信息、recordId", "提交测验并保存历史结果"),
        ("学习助手", "/assistant/chat", "POST", "普通用户", "message、sessionId、context", "模型回复与会话补全结果", "基于上下文发起学习问答"),
        ("学习看板", "/dashboard/overview", "GET", "普通用户", "可选时间范围参数", "学习概览统计数据", "展示今日学习、复习和测验统计"),
        ("后台词库", "/admin/word-bank/import", "POST", "管理员", "批次信息、词库文件内容", "导入数量、跳过数量、批次号", "执行词库导入并生成批次记录"),
        ("后台词库", "/admin/word-bank/rollback", "POST", "管理员", "batchId", "回滚结果与批次状态", "按批次撤销导入操作"),
        ("模型管理", "/admin/llm/cloud-models", "GET/POST", "管理员", "模型查询参数或模型配置字段", "模型列表或新增结果", "查询或维护云模型配置"),
    ]
    add_table_title_and_table_before(doc, anchor_45, "表4-5 系统核心接口设计表", ["模块", "接口路径", "方法", "权限", "请求参数", "返回结果", "功能说明"], interface_rows, [1.8, 3.5, 1.5, 1.8, 4.0, 4.0, 3.6])
    add_text_before(anchor_45, "表4-5中的接口可以看出，普通用户侧接口以学习闭环为中心，管理员侧接口则更多承担词库和模型治理职责。统一结果封装、统一路径风格以及控制层职责分离，使前端在调用这些接口时能够保持较高的一致性，也降低了功能扩展和联调成本。")

    add_text_before(anchor_45, "4.4.2 顺序图", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(anchor_45, "单词详情页触发 AI 内容生成的时序关系如图4-4所示。该顺序图围绕用户、前端详情页、WordController、WordServiceImpl、Redis、MySQL、模型适配解析服务以及本地/云端模型展开，用于说明请求如何从页面进入后台，再完成结果解析、回填与刷新。")
    add_caption_before(anchor_45, "（此处插入图4-4 顺序图）", placeholder=True)
    add_caption_before(anchor_45, "图4-4 单词详情页 AI 内容生成顺序图")
    add_text_before(anchor_45, "由图4-4可见，系统首先通过 `/word/detail` 读取缓存和数据库中的稳定内容，在页面确认仍有缺失字段后，再通过 `/word/llm/generate` 触发模型调用。模型返回结果后，服务端并不会直接原样输出，而是要经过提示解析、结构校验、`word_meta` 回填和 Redis 缓存刷新等步骤，最后才由页面轮询或重试机制获取最终结果。这种顺序安排保证了生成链路既能利用缓存，又能保留失败降级和异步回填能力。")
    add_text_before(anchor_45, "为保证新增图表和表格的版式一致性，本文统一采用“图题置于图下方、表题置于表上方”的排版方式，图题和表题均居中排列，编号按章节连续顺延。图宽控制在版心范围内，优先采用横向布局，图中节点文字尽量控制在两行以内，连线不穿越文字区域；表格以三线表风格为优先，若字段较多则保留网格线，但表头必须加粗，且表内英文、数字、接口路径与数据表名均使用 Times New Roman 字体表示。")


def use_case_drawio():
    cells = []
    cells.append(cell("10", "四六级高频词记忆系统", "rounded=1;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;fontColor=#222222;fontStyle=1;", True, geometry=mx_geometry(260, 70, 980, 540)))
    actor_style = "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;strokeColor=#444444;fontColor=#222222;"
    usecase_style = "ellipse;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;fontColor=#222222;"
    edge_style = "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=18;html=1;endArrow=none;strokeColor=#444444;entryX=0.5;entryY=0.5;exitX=0.5;exitY=0.5;"
    cells.extend([
        cell("11", "普通用户", actor_style, True, geometry=mx_geometry(100, 170, 70, 120)),
        cell("12", "管理员", actor_style, True, geometry=mx_geometry(100, 410, 70, 120)),
        cell("13", "本地模型服务", actor_style, True, geometry=mx_geometry(1320, 170, 90, 120)),
        cell("14", "云端模型服务", actor_style, True, geometry=mx_geometry(1320, 410, 90, 120)),
    ])
    for cid, label, x, y in [
        ("21", "注册登录", 350, 130), ("22", "词库浏览", 570, 130), ("23", "查看单词详情", 790, 130), ("24", "加入学习计划", 1010, 130),
        ("25", "AI 内容生成", 350, 285), ("26", "完成每日复习", 570, 285), ("27", "发起模拟测验", 790, 285), ("28", "学习助手对话", 1010, 285),
        ("29", "看板统计查看", 460, 445), ("30", "词库导入回滚", 690, 445), ("31", "模型管理", 920, 445),
    ]:
        cells.append(cell(cid, label, usecase_style, True, geometry=mx_geometry(x, y, 145, 56)))
    for eid, src, dst in [
        ("101", "11", "21"), ("102", "11", "22"), ("103", "11", "23"), ("104", "11", "24"),
        ("105", "11", "25"), ("106", "11", "26"), ("107", "11", "27"), ("108", "11", "28"), ("109", "11", "29"),
        ("110", "12", "30"), ("111", "12", "31"),
        ("112", "13", "25"), ("113", "13", "28"), ("114", "14", "25"), ("115", "14", "28"),
    ]:
        cells.append(cell(eid, "", edge_style, edge=True, source=src, target=dst, geometry=mx_geometry(0, 0, 0, 0, relative=True)))
    return drawio_file("图3-1-系统用例图", cells)


def dfd_drawio():
    cells = []
    entity_style = "rounded=0;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;fontColor=#222222;"
    process_style = "rounded=1;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;fontColor=#222222;"
    store_style = "shape=partialRectangle;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;right=0;fontColor=#222222;"
    edge_style = "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=18;html=1;endArrow=block;endFill=1;strokeColor=#444444;"
    cells.extend([
        cell("11", "用户", entity_style, True, geometry=mx_geometry(70, 180, 100, 56)),
        cell("12", "管理员", entity_style, True, geometry=mx_geometry(70, 420, 100, 56)),
        cell("13", "本地模型", entity_style, True, geometry=mx_geometry(1400, 210, 120, 56)),
        cell("14", "云端模型", entity_style, True, geometry=mx_geometry(1400, 410, 120, 56)),
        cell("21", "认证授权", process_style, True, geometry=mx_geometry(300, 110, 145, 56)),
        cell("22", "词库学习", process_style, True, geometry=mx_geometry(540, 110, 145, 56)),
        cell("23", "AI 内容生成", process_style, True, geometry=mx_geometry(780, 110, 145, 56)),
        cell("24", "复习调度", process_style, True, geometry=mx_geometry(300, 320, 145, 56)),
        cell("25", "测验管理", process_style, True, geometry=mx_geometry(540, 320, 145, 56)),
        cell("26", "学习助手", process_style, True, geometry=mx_geometry(780, 320, 145, 56)),
        cell("27", "后台词库管理", process_style, True, geometry=mx_geometry(1020, 320, 165, 56)),
        cell("31", "user", store_style, True, geometry=mx_geometry(1030, 90, 150, 64)),
        cell("32", "词表 cet4zx/cet6zx", store_style, True, geometry=mx_geometry(1210, 90, 180, 64)),
        cell("33", "word_meta", store_style, True, geometry=mx_geometry(1030, 190, 150, 64)),
        cell("34", "user_word_progress", store_style, True, geometry=mx_geometry(1210, 190, 180, 64)),
        cell("35", "review_log", store_style, True, geometry=mx_geometry(1030, 400, 150, 64)),
        cell("36", "quiz_session_record", store_style, True, geometry=mx_geometry(1210, 400, 180, 64)),
        cell("37", "word_import_batch", store_style, True, geometry=mx_geometry(1030, 500, 150, 64)),
        cell("38", "cloud_llm_model", store_style, True, geometry=mx_geometry(1210, 500, 180, 64)),
    ])
    for eid, src, dst, label in [
        ("101", "11", "21", "注册/登录"), ("102", "11", "22", "词库查询"), ("103", "11", "23", "详情与生成"),
        ("104", "11", "24", "复习评分"), ("105", "11", "25", "测验作答"), ("106", "11", "26", "助手问答"),
        ("107", "12", "27", "导入/回滚"), ("108", "12", "23", "模型配置维护"),
        ("109", "21", "31", "用户信息"), ("110", "22", "32", "词汇读取"), ("111", "23", "33", "生成回填"),
        ("112", "23", "38", "读取模型配置"), ("113", "24", "34", "进度更新"), ("114", "24", "35", "日志写入"),
        ("115", "25", "36", "测验记录"), ("116", "26", "33", "对话上下文"), ("117", "27", "37", "批次记录"),
        ("118", "23", "13", "本地调用"), ("119", "23", "14", "云端调用"), ("120", "26", "13", "本地问答"),
        ("121", "26", "14", "云端问答"),
    ]:
        cells.append(cell(eid, label, edge_style, edge=True, source=src, target=dst, geometry=mx_geometry(0, 0, 0, 0, relative=True)))
    return drawio_file("图3-2-数据流图", cells)


def activity_drawio():
    cells = []
    start_style = "ellipse;whiteSpace=wrap;html=1;aspect=fixed;strokeColor=#444444;fillColor=#ffffff;"
    process_style = "rounded=1;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;"
    decision_style = "rhombus;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;"
    edge_style = "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=18;html=1;endArrow=block;endFill=1;strokeColor=#444444;"
    cells.extend([
        cell("11", "开始", start_style, True, geometry=mx_geometry(90, 250, 50, 50)),
        cell("12", "进入详情页", process_style, True, geometry=mx_geometry(180, 245, 150, 56)),
        cell("13", "请求基础词条信息", process_style, True, geometry=mx_geometry(390, 245, 170, 56)),
        cell("14", "检查缓存", process_style, True, geometry=mx_geometry(620, 245, 130, 56)),
        cell("15", "缓存命中？", decision_style, True, geometry=mx_geometry(810, 228, 120, 90)),
        cell("16", "直接展示详情", process_style, True, geometry=mx_geometry(1000, 120, 150, 56)),
        cell("17", "查询数据库/元数据", process_style, True, geometry=mx_geometry(1000, 360, 170, 56)),
        cell("18", "内容缺失？", decision_style, True, geometry=mx_geometry(1230, 343, 120, 90)),
        cell("19", "触发模型生成", process_style, True, geometry=mx_geometry(1430, 360, 150, 56)),
        cell("20", "解析并校验结果", process_style, True, geometry=mx_geometry(1640, 360, 160, 56)),
        cell("21", "回填数据库与缓存", process_style, True, geometry=mx_geometry(1860, 360, 170, 56)),
        cell("22", "返回详情页展示", process_style, True, geometry=mx_geometry(1430, 120, 170, 56)),
        cell("23", "返回降级结果", process_style, True, geometry=mx_geometry(1640, 500, 160, 56)),
        cell("24", "结束", start_style, True, geometry=mx_geometry(1670, 120, 50, 50)),
    ])
    for eid, src, dst, label in [
        ("101", "11", "12", ""), ("102", "12", "13", ""), ("103", "13", "14", ""), ("104", "14", "15", ""),
        ("105", "15", "16", "是"), ("106", "15", "17", "否"), ("107", "17", "18", ""), ("108", "18", "22", "否"),
        ("109", "18", "19", "是"), ("110", "19", "20", ""), ("111", "20", "21", "成功"), ("112", "20", "23", "异常"),
        ("113", "21", "22", ""), ("114", "16", "24", ""), ("115", "22", "24", ""), ("116", "23", "24", ""),
    ]:
        cells.append(cell(eid, label, edge_style, edge=True, source=src, target=dst, geometry=mx_geometry(0, 0, 0, 0, relative=True)))
    return drawio_file("图3-3-活动图", cells)


def sequence_drawio():
    cells = []
    header_style = "rounded=0;whiteSpace=wrap;html=1;strokeColor=#444444;fillColor=#ffffff;fontStyle=1;align=center;verticalAlign=middle;"
    lifeline_style = "shape=line;strokeColor=#666666;dashed=1;direction=south;"
    msg_style = "rounded=0;html=1;endArrow=block;endFill=1;strokeColor=#444444;fontColor=#222222;labelBackgroundColor=#ffffff;labelBorderColor=none;"
    return_style = "rounded=0;html=1;endArrow=open;dashed=1;strokeColor=#555555;fontColor=#222222;labelBackgroundColor=#ffffff;labelBorderColor=none;"
    participants = [
        ("11", "用户", 110, 120),
        ("12", "前端详情页", 320, 150),
        ("13", "WordController", 560, 170),
        ("14", "WordServiceImpl", 820, 190),
        ("15", "Redis", 1080, 120),
        ("16", "MySQL", 1280, 120),
        ("17", "模型适配/解析服务", 1510, 190),
        ("18", "本地/云端模型", 1780, 160),
    ]
    x_map = {}
    for cid, label, x, w in participants:
        cells.append(cell(cid, label, header_style, True, geometry=mx_geometry(x, 40, w, 50)))
        line_x = x + (w / 2)
        x_map[cid] = line_x
        cells.append(cell(f"{cid}l", "", lifeline_style, True, geometry=mx_geometry(line_x, 100, 1, 780)))

    messages = [
        ("101", "11", "12", "打开详情页", 140, msg_style),
        ("102", "12", "13", "请求 /word/detail", 190, msg_style),
        ("103", "13", "14", "查询详情", 240, msg_style),
        ("104", "14", "15", "读取缓存", 290, msg_style),
        ("105", "15", "14", "缓存命中/未命中", 340, return_style),
        ("106", "14", "16", "读取词表与元数据", 390, msg_style),
        ("107", "16", "14", "返回基础数据", 440, return_style),
        ("108", "12", "13", "触发 /word/llm/generate", 520, msg_style),
        ("109", "13", "14", "组装生成请求", 570, msg_style),
        ("110", "14", "17", "构造提示词", 620, msg_style),
        ("111", "17", "18", "调用本地/云端模型", 670, msg_style),
        ("112", "18", "17", "返回结构化结果", 720, return_style),
        ("113", "17", "14", "解析并校验结果", 770, return_style),
        ("114", "14", "16", "回填 word_meta", 820, msg_style),
        ("115", "14", "15", "刷新缓存", 870, msg_style),
        ("116", "12", "13", "轮询/重试详情", 950, msg_style),
        ("117", "13", "14", "重新读取缓存/数据库", 1000, msg_style),
        ("118", "14", "12", "返回 pending/complete", 1050, return_style),
        ("119", "12", "11", "页面展示结果", 1100, return_style),
    ]
    for eid, src, dst, label, y, style in messages:
        cells.append(point_edge(eid, label, style, x_map[src], y, x_map[dst], y))
    return drawio_file("图4-4-顺序图", cells)


def generate_diagrams():
    REPO_DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    DOWNLOAD_DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    outputs = {
        "图3-1-系统用例图.drawio": use_case_drawio(),
        "图3-2-数据流图.drawio": dfd_drawio(),
        "图3-3-活动图.drawio": activity_drawio(),
        "图4-4-顺序图.drawio": sequence_drawio(),
    }
    for filename, content in outputs.items():
        write_text(REPO_DIAGRAM_DIR / filename, content)
        write_text(DOWNLOAD_DIAGRAM_DIR / filename, content)


def update_docx_and_pdf():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not BACKUP_PATH.exists():
        shutil.copyfile(DOCX_PATH, BACKUP_PATH)
    doc = Document(DOCX_PATH)
    insert_section_34(doc)
    insert_section_43_and_44(doc)
    doc.save(DOCX_PATH)

    import win32com.client
    from win32com.client import constants

    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        com_doc = word.Documents.Open(str(DOCX_PATH))
        if com_doc.TablesOfContents.Count >= 1:
            com_doc.TablesOfContents(1).Update()
        com_doc.Fields.Update()
        com_doc.Save()
        com_doc.SaveAs(str(PDF_PATH), FileFormat=constants.wdFormatPDF)
        com_doc.Close(False)
    finally:
        word.Quit()


def main():
    generate_diagrams()
    update_docx_and_pdf()
    print(f"DOCX={DOCX_PATH}")
    print(f"PDF={PDF_PATH}")
    print(f"BACKUP={BACKUP_PATH}")
    for name in ["图3-1-系统用例图.drawio", "图3-2-数据流图.drawio", "图3-3-活动图.drawio", "图4-4-顺序图.drawio"]:
        print(f"DIAGRAM={DOWNLOAD_DIAGRAM_DIR / name}")


if __name__ == "__main__":
    main()
