from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_DOCX = Path(r"C:/Users/zjf20/Desktop/毕业设计（论文）-项目重写稿.docx")
OUTPUT_DOCX = Path(r"D:/zjf20/Downloads/毕业设计（论文）-项目重写稿.docx")
ROOT = Path(r"D:/JAVA/ideaProjects/cet46-vocab")

ASSET_PLAN = [
    ("4.1 系统总体架构设计", ROOT / "docs/thesis-assets/diagrams/system-architecture.png", "图4-1 系统总体架构图", 15.5),
    ("4.2 功能模块设计", ROOT / "docs/thesis-assets/diagrams/functional-modules.png", "图4-2 系统功能模块图", 15.5),
    ("4.3 数据库设计", ROOT / "docs/thesis-assets/diagrams/er-diagram.png", "图4-3 系统 E-R 图", 15.5),
    ("5.1 用户认证与权限控制实现", ROOT / "docs/thesis-assets/screenshots/system-onboarding.png", "图5-1 用户注册与学习风格初始化页面", 14.5),
    ("5.2 词库导入与词条管理实现", ROOT / "docs/thesis-assets/screenshots/system-word-list.png", "图5-2 词库浏览与加入学习页面", 15.0),
    ("5.3 单词详情页与 AI 生成链路实现", ROOT / "docs/thesis-assets/screenshots/system-word-detail.png", "图5-3 单词详情与 AI 解释生成页面", 15.0),
    ("5.4 学习助手与本地/云端双模型适配实现", ROOT / "docs/thesis-assets/screenshots/system-assistant.png", "图5-4 学习助手页面运行效果", 15.0),
    ("5.5 复习调度与学习看板实现", ROOT / "docs/thesis-assets/screenshots/system-dashboard.png", "图5-5 学习看板页面运行效果", 15.0),
    ("5.6 模拟测验与历史记录实现", ROOT / "docs/thesis-assets/screenshots/system-quiz.png", "图5-6 模拟测验页面运行效果", 15.0),
]


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.add_paragraph()._element  # placeholder


def paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.paragraphs[
        paragraph._parent._element.index(new_p)  # type: ignore[attr-defined]
    ]


def set_run_fonts(run, size_pt=10.5, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def format_caption(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=10.5)


def format_picture_paragraph(paragraph, image_path, width_cm):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def build_paragraph_after(paragraph):
    new_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_xml)
    return paragraph._parent.paragraphs[-1]


def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph


def insert_asset_after_heading(document, heading_text, image_path, caption, width_cm):
    paragraphs = list(iter_paragraphs(document))
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading_text:
            continue

        anchor = paragraph
        pic_para = document.add_paragraph()
        pic_xml = pic_para._p
        pic_para._p.getparent().remove(pic_xml)
        anchor._p.addnext(pic_xml)
        format_picture_paragraph(pic_para, image_path, width_cm)

        cap_para = document.add_paragraph()
        cap_xml = cap_para._p
        cap_para._p.getparent().remove(cap_xml)
        pic_para._p.addnext(cap_xml)
        format_caption(cap_para, caption)
        return True
    return False


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Missing source docx: {SOURCE_DOCX}")

    document = Document(str(SOURCE_DOCX))

    missing = []
    for heading_text, image_path, caption, width_cm in ASSET_PLAN:
        if not image_path.exists():
            missing.append(str(image_path))
            continue
        inserted = insert_asset_after_heading(document, heading_text, image_path, caption, width_cm)
        if not inserted:
            missing.append(f"Heading not found: {heading_text}")

    if missing:
        raise RuntimeError("Asset insertion failed for:\n" + "\n".join(missing))

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
