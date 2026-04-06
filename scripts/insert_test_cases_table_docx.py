from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")

TEST_CASES = [
    ("TC-01", "用户注册与风格初始化", "进入注册页，填写用户名与密码并完成三道风格题", "用户成功注册并跳转到初始化页面，风格偏好写入用户资料"),
    ("TC-02", "词库浏览与加入学习", "已登录普通用户，进入词库列表页并筛选单词", "系统可按词库和词性筛选单词，并将目标单词加入学习计划"),
    ("TC-03", "单词详情 AI 内容生成", "从词库列表进入单词详情页，等待异步生成完成", "系统先展示基础释义，再补齐例句、助记和扩展解释"),
    ("TC-04", "复习任务读取与评分提交", "存在待复习单词，进入复习页并完成评分", "系统正确更新下次复习时间、掌握状态和学习统计"),
    ("TC-05", "模拟测验提交与历史查看", "进入测验页后选择题量、模式和词库", "提交后返回得分与错题信息，历史记录可查看当次详情"),
    ("TC-06", "学习助手问答", "登录后进入学习助手页并输入问题", "系统返回与当前问题相关的解释或建议，异常时可回退到可用模型"),
]


def set_run_fonts(run, size_pt=10.5, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_paragraph_text(paragraph, text, size_pt=10.5, bold=False, align=None, first_line_indent=True):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)


def insert_paragraph_before(paragraph, text, size_pt=10.5, align=None, first_line_indent=True):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    set_paragraph_text(new_para, text, size_pt=size_pt, align=align, first_line_indent=first_line_indent)
    return new_para


def format_cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(18)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_run_fonts(run, size_pt=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    doc = Document(str(DOCX_PATH))

    for p in doc.paragraphs:
        if p.text.strip() == "表6-1 典型功能测试用例":
            print("table already exists")
            return

    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "6.4 部署过程与运行支撑":
            target = p
            break
    if target is None:
        raise RuntimeError("Target heading not found")

    insert_paragraph_before(
        target,
        "为使测试章节具备更明确的可复查性，结合当前仓库中的测试说明文档、自动化回归结果和实际运行页面，选取具有代表性的功能测试用例如表6-1所示。",
        size_pt=12,
        first_line_indent=True,
    )
    insert_paragraph_before(
        target,
        "表6-1 典型功能测试用例",
        size_pt=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )

    table = doc.add_table(rows=len(TEST_CASES) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["编号", "测试内容", "测试步骤", "预期结果"]
    for col, title in enumerate(headers):
        format_cell(table.cell(0, col), title, bold=True)

    for row_idx, row in enumerate(TEST_CASES, start=1):
        for col_idx, value in enumerate(row):
            format_cell(table.cell(row_idx, col_idx), value)

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    target._p.addprevious(tbl)

    insert_paragraph_before(
        target,
        "从表6-1可以看出，测试章节不仅给出了自动化回归与构建结果，还补充了围绕注册、词库浏览、AI 生成、复习、测验和学习助手等核心业务流程的典型功能测试项，能够更完整地支撑论文对系统可用性的说明。",
        size_pt=12,
        first_line_indent=True,
    )

    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
