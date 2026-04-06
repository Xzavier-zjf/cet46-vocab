from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")

COVER_ROWS = [
    ("学    院", "计算机工程学院/大数据学院"),
    ("专业班级", "2022软件工程4班"),
    ("学生姓名", ""),
    ("学生学号", ""),
    ("指导教师", ""),
    ("提交日期", "年    月    日"),
]

FIGURES = [
    ("图4-1 系统总体架构图", Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/diagrams/system-architecture.png"), 15.5),
    ("图4-2 系统功能模块图", Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/diagrams/functional-modules.png"), 15.5),
    ("图4-3 系统 E-R 图", Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/diagrams/er-diagram.png"), 15.5),
    ("图6-3 系统部署流程图", Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/diagrams/deployment-flow.png"), 14.5),
]


def set_run_fonts(run, size_pt=12, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def insert_table_before(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(13))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addprevious(tbl)
    return table


def insert_picture_before(paragraph, image_path: Path, width_cm: float):
    pic_para = paragraph._parent.add_paragraph()
    pic_xml = pic_para._p
    pic_xml.getparent().remove(pic_xml)
    paragraph._p.addprevious(pic_xml)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.first_line_indent = Pt(0)
    pic_para.paragraph_format.line_spacing = Pt(20)
    pic_para.paragraph_format.space_before = Pt(0)
    pic_para.paragraph_format.space_after = Pt(0)
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return pic_para


def format_cover_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, (label, value) in enumerate(COVER_ROWS):
        left = table.cell(row_idx, 0)
        right = table.cell(row_idx, 1)
        left.width = Cm(4.2)
        right.width = Cm(8.8)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for cell, text in [(left, label), (right, value)]:
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(text)
            set_run_fonts(run, size_pt=12)

        tc_pr = right._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")

        # remove other borders
        for tag in ["top", "left", "right", "insideH", "insideV"]:
            node = borders.find(qn(f"w:{tag}"))
            if node is None:
                node = OxmlElement(f"w:{tag}")
                borders.append(node)
            node.set(qn("w:val"), "nil")

    for row in table.rows:
        row.height = Cm(0.82)


def rebuild_cover(doc):
    anchor = doc.paragraphs[11]
    table = insert_table_before(anchor, len(COVER_ROWS), 2)
    format_cover_table(table)
    for idx in range(11, 17):
        clear_paragraph(doc.paragraphs[idx])


def replace_target_figures(doc):
    for caption, image_path, width_cm in FIGURES:
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        paragraphs = list(doc.paragraphs)
        target_idx = None
        for idx, p in enumerate(paragraphs):
            if p.text.strip() == caption:
                target_idx = idx
                break
        if target_idx is None:
            raise RuntimeError(f"Caption not found: {caption}")

        # remove immediate empty paragraphs before caption; these usually hold the old image container/spacers
        idx = target_idx - 1
        removed = 0
        while idx >= 0 and removed < 4:
            if paragraphs[idx].text.strip():
                break
            delete_paragraph(paragraphs[idx])
            removed += 1
            idx -= 1
        paragraphs = list(doc.paragraphs)
        for p in paragraphs:
            if p.text.strip() == caption:
                insert_picture_before(p, image_path, width_cm)
                break


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    doc = Document(str(DOCX_PATH))
    rebuild_cover(doc)
    replace_target_figures(doc)
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
