from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


DOCX_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-答辩前清稿.docx")


def set_run_fonts(run, size_pt=12, font_cn="宋体", font_en="Times New Roman"):
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_text(paragraph, text):
    clear_paragraph(paragraph)
    if text:
        run = paragraph.add_run(text)
        set_run_fonts(run, size_pt=12)
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if text else Pt(0)


def main():
    doc = Document(str(DOCX_PATH))
    replacements = {
        217: "",
        228: "",
        295: "",
        236: "由图5-3可见，单词详情页并不等待完整生成结束后再返回页面，而是先返回基础数据，再通过异步生成与状态回填逐步补齐扩展内容。随后，页面会按轮询结果刷新扩展字段，其实际运行效果如图5-4所示。",
    }
    for idx, text in replacements.items():
        if idx < len(doc.paragraphs):
            set_text(doc.paragraphs[idx], text)
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
