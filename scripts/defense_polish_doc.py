from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_DOCX = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")
OUTPUT_DOCX = Path(r"D:/zjf20/Downloads/毕业设计（论文）-答辩前清稿.docx")


PARAGRAPH_REWRITES = {
    "结合任务书、开题报告和当前仓库实现，本文研究目标可归纳为四个方面。": "结合任务书、开题报告以及当前仓库已经落地的代码实现，本文的研究目标可以归纳为四个方面。首先，围绕 CET-4 与 CET-6 高频词汇建立较为规范的数据组织方式，并通过批量导入与批次回滚机制支撑词库的持续维护。其次，面向单词详情页和学习助手构建可复用的 AI 解释生成链路，使系统能够稳定输出例句、近义词辨析、助记内容和扩展解释，同时兼容本地模型与云端模型两种运行路径。",
    "第三，将 SM-2 间隔重复算法嵌入学习进度管理过程，根据用户在复习环节中的评分动态更新复习间隔、重复次数和下一次复习日期，进而形成“加入学习—进入复习队列—记录学习状态—在看板中反馈结果”的闭环。第四，在学习功能之外补充模拟测验、学习看板、模型使用统计和模型配置管理等配套模块，使系统不仅能生成学习内容，还能支撑阶段性训练和学习行为分析。": "再次，将 SM-2 间隔重复算法嵌入学习进度管理过程，根据用户在复习环节中的评分实时更新复习间隔、重复次数和下一次复习日期，从而把“加入学习、进入复习、提交评分、统计反馈”连接成完整闭环。最后，在学习功能之外进一步补充模拟测验、学习看板、模型使用统计和模型配置管理等模块，使系统不只提供内容生成能力，还能承担阶段性训练与学习行为分析的任务。",
    "从整体设计看，系统沿用浏览器/服务器模式展开。": "结合图4-1可以看出，系统总体结构沿用浏览器/服务器模式展开。",
    "表示层主要包括 Dashboard、WordList、WordDetail、Review、Quiz、Assistant、Profile 和管理端页面。": "进一步观察图4-1可以发现，表示层主要包括 Dashboard、WordList、WordDetail、Review、Quiz、Assistant、Profile 和管理端页面。",
    "该模块负责注册、登录、登出、资料查看和角色边界控制。": "由图4-2可见，用户与权限模块负责注册、登录、登出、资料查看以及角色边界控制。",
    "根据当前数据库脚本和迁移文件，系统数据库由词表、用户表、进度表、日志表、导入批次表和模型相关表共同构成。": "结合图4-3可见，系统数据库由词表、用户表、进度表、日志表、导入批次表和模型相关表共同构成。",
    "用户认证部分围绕 AuthController、AuthServiceImpl 与前端路由守卫协同工作。": "如图5-1所示，用户完成注册后需要先进入学习风格初始化页面，再由前后端协同完成偏好设置与后续导航。对应到实现层面，用户认证部分围绕 AuthController、AuthServiceImpl 与前端路由守卫共同工作。",
    "词库导入模块主要由管理端接口完成。": "如图5-2所示，普通用户在词库列表页可以直接浏览词条、筛选词性并把目标单词加入学习计划；而在管理端，词库导入模块主要通过后台接口完成。",
    "单词详情页是系统最能体现“大模型解释能力”的页面之一。": "结合图5-3和图5-4可以更直观地看到，单词详情页既是 AI 生成链路的触发入口，也是系统最能体现“大模型解释能力”的页面之一。",
    "学习助手模块用于承接词条详情之外的延伸学习需求。": "如图5-5所示，学习助手模块用于承接词条详情之外的延伸学习需求。",
    "复习模块前端采用翻卡式交互。": "如图5-6所示，复习与看板模块在页面层面形成了“任务提醒 + 数据反馈”的组合关系。具体到复习模块，前端采用翻卡式交互。",
    "测验模块的设计目标是为用户提供阶段性检验能力。": "如图5-7所示，模拟测验页面支持在同一入口中完成题量选择、模式切换、作答与结果查看。对应的后端实现中，测验模块的设计目标是为用户提供阶段性检验能力。",
    "自动化测试部分以当前仓库中能够稳定复现的测试类为依据，重点覆盖了 SM-2 算法、用户控制器、用户服务、云模型边界和角色权限等模块。": "如图6-1所示，当前论文中采用的测试证据并非抽象描述，而是直接来自仓库中的测试报告、构建结果与执行记录。自动化测试部分以当前仓库中能够稳定复现的测试类为依据，重点覆盖了 SM-2 算法、用户控制器、用户服务、云模型边界和角色权限等模块。",
    "根据仓库中的 `deploy.sh` 脚本，系统部署流程包括拉取最新代码、前端构建、后端打包、停止旧进程、启动新 Jar 和重载 Nginx 六个步骤。": "结合图6-2与图6-3可以看出，当前项目已经具备较清晰的部署支撑材料。根据仓库中的 `deploy.sh` 脚本，系统部署流程包括拉取最新代码、前端构建、后端打包、停止旧进程、启动新 Jar 和重载 Nginx 六个步骤。",
    "本文围绕当前代码仓库中的 CET46 Vocabulary 项目，对一个“以词汇学习为主线、以大模型解释为增强、以复习调度为核心”的系统进行了梳理与总结。": "本文围绕当前代码仓库中的 CET46 Vocabulary 项目，对一个“以词汇学习为主线、以大模型解释为增强、以复习调度为核心”的系统进行了梳理、分析与总结。",
}


FIGURE_INTROS = {
    "图4-2 系统功能模块图": "系统功能模块划分情况如图4-2所示。",
    "图4-3 系统 E-R 图": "数据库核心实体及其关联关系如图4-3所示。",
    "图5-1 用户注册与学习风格初始化页面": "用户注册及学习风格初始化页面的运行效果如图5-1所示。",
    "图5-2 词库浏览与加入学习页面": "词库浏览与加入学习页面的运行效果如图5-2所示。",
    "图5-3 单词详情 AI 生成与回填时序图": "单词详情页触发 AI 生成、缓存回填和前端轮询展示的关键交互过程如图5-3所示。",
    "图6-2 部署脚本与启动日志证据截图": "部署脚本、Nginx 配置和后端启动日志整理后的证据截图如图6-2所示。",
}


FIGURE_POSTS = {
    "图4-2 系统功能模块图": "图4-2从模块边界角度概括了系统的主要组成部分，为后续分模块展开实现分析提供了结构参照。",
    "图4-3 系统 E-R 图": "图4-3说明数据库设计并非围绕单一词表展开，而是围绕词汇、用户、进度、日志和导入批次等对象形成协同关系。",
    "图5-3 单词详情 AI 生成与回填时序图": "由图5-3可见，单词详情页并不等待完整生成结束后再返回页面，而是先返回基础数据，再通过异步生成与状态回填逐步补齐扩展内容。",
    "图6-2 部署脚本与启动日志证据截图": "图6-2从工程证据角度说明，论文第六章关于部署过程的描述与仓库中的实际脚本和运行日志保持一致。",
    "图6-3 系统部署流程图": "图6-3将脚本中的关键步骤进一步抽象为部署流程，便于从论文层面概括系统上线时的执行顺序与检查节点。",
}


DECLARATION_TEXT = "本人郑重声明：所呈交的毕业设计（论文）是在指导教师指导下独立完成的研究成果。除文中已经明确标注引用的内容外，论文不包含任何他人或集体已经公开发表、撰写的研究成果。凡对本研究有重要帮助的个人和集体，均已在文中以适当方式注明。本人愿对本声明所引起的一切后果承担相应责任。"
LICENSE_TEXT = "本人完全了解广州城市理工学院关于毕业设计（论文）保存、使用与管理的有关规定，同意学校按照要求保存本论文的纸质版和电子版，并在校内教学、科研与文献服务范围内进行检索、阅览、复制或数字化保存。在不以营利为目的的前提下，学校可以依法公开论文的部分或全部内容。"

ACK_TEXTS = [
    "毕业设计（论文）完成过程中，指导教师在选题论证、研究思路梳理、系统实现分析和论文结构修改等方面给予了细致指导与持续帮助。在论文撰写和项目整理过程中，这些建议使本文能够从单纯的代码说明进一步转化为较为完整的工程总结，在此谨致以诚挚谢意。",
    "同时，感谢学院提供的学习环境与实践条件，感谢大学期间各位任课教师在软件工程、数据库原理、Web 开发、测试方法和人工智能应用等课程中打下的基础。正是这些课程训练，使本文在分析系统需求、数据库结构、关键模块与测试部署时能够建立在较扎实的专业知识之上。",
    "此外，感谢在项目联调、资料整理和阶段性修改过程中给予帮助与建议的同学和朋友。由于学校最终提交材料仍包含签名、日期及其他个人信息项，相关内容将由本人在提交前按照学院要求补充完善。后续本人也将结合导师意见和答辩反馈，继续对细节表述、图表排版和个别章节文字做进一步打磨。"
]


def set_run_fonts(run, size_pt=12, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_paragraph_text(paragraph, text, size_pt=12, align=None, first_line_indent=True):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    return new_para


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    return new_para


def prev_nonempty(paragraphs, idx):
    for i in range(idx - 1, -1, -1):
        if paragraphs[i].text.strip():
            return i, paragraphs[i]
    return None, None


def next_nonempty(paragraphs, idx):
    for i in range(idx + 1, len(paragraphs)):
        if paragraphs[i].text.strip():
            return i, paragraphs[i]
    return None, None


def apply_rewrites(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        for prefix, replacement in PARAGRAPH_REWRITES.items():
            if text.startswith(prefix):
                set_paragraph_text(p, replacement, size_pt=12, first_line_indent=True)
                break


def polish_front_and_back(doc):
    # fixed positions inherited from template
    set_paragraph_text(doc.paragraphs[19], DECLARATION_TEXT, size_pt=10.5, first_line_indent=True)
    set_paragraph_text(doc.paragraphs[25], LICENSE_TEXT, size_pt=10.5, first_line_indent=True)
    ack_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "致  谢":
            ack_idx = i
            break
    if ack_idx is not None:
        for offset, text in enumerate(ACK_TEXTS, start=1):
            idx = ack_idx + offset
            if idx < len(doc.paragraphs):
                set_paragraph_text(doc.paragraphs[idx], text, size_pt=10.5, first_line_indent=True)


def ensure_figure_mentions(doc):
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        caption = p.text.strip()
        if caption not in FIGURE_INTROS:
            continue
        prev_idx, prev_p = prev_nonempty(paragraphs, idx)
        intro = FIGURE_INTROS[caption]
        if prev_p is None or prev_p.text.strip() != intro:
            new_p = insert_paragraph_before(p)
            set_paragraph_text(new_p, intro, size_pt=12, first_line_indent=True)
            paragraphs = doc.paragraphs

    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        caption = p.text.strip()
        if caption not in FIGURE_POSTS:
            continue
        next_idx, next_p = next_nonempty(paragraphs, idx)
        post = FIGURE_POSTS[caption]
        if next_p is None or next_p.style.name.startswith("Heading"):
            new_p = insert_paragraph_after(p)
            set_paragraph_text(new_p, post, size_pt=12, first_line_indent=True)
        elif not next_p.text.strip().startswith(("由图", "结合图", "图4-", "图5-", "图6-")):
            set_paragraph_text(next_p, post + next_p.text.strip(), size_pt=12, first_line_indent=True)


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    doc = Document(str(SOURCE_DOCX))
    apply_rewrites(doc)
    polish_front_and_back(doc)
    ensure_figure_mentions(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
