from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_DOCX = Path(r"C:/Users/zjf20/Desktop/毕业设计（论文）-项目重写稿.docx")
OUTPUT_DOCX = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")
ROOT = Path(r"D:/JAVA/ideaProjects/cet46-vocab")


TEXT_REPLACEMENTS = {
    "本文围绕“集成大模型解释能力的四六级高频词记忆系统”展开设计与实现": "本文以当前仓库中的 CET46 Vocabulary 项目为基础，对“集成大模型解释能力的四六级高频词记忆系统”进行了系统化梳理和实现总结。论文关注的不是单一的背词页面，而是一个围绕词汇学习闭环展开的完整系统：前端负责词库浏览、单词详情、复习、测验、学习助手与统计看板等交互页面，后端负责认证授权、词汇数据管理、AI 内容生成、学习记录维护、测验结果存储以及模型配置管理。与任务书中偏向“轻量化背词工具”的设想相比，现有项目已经进一步扩展出本地 Ollama 与云端 OpenAI 兼容接口双通道接入、异步生成与状态回填、私有模型配置、批量导入回滚、学习看板统计等工程能力。全文将围绕该项目的需求、总体结构、关键模块、测试证据和部署方式展开说明，并对系统当前可达到的效果与后续改进空间进行归纳。",
    "大学英语四、六级考试仍然是高校英语学习中的重要评价节点": "在多数高校的英语学习场景中，四六级备考依旧是学生接触频率最高、投入时间最长的任务之一。和阅读、听力这类综合训练相比，词汇记忆更适合被拆解成高频、小步、可追踪的学习过程，因此也最容易落地为软件系统。现实中常见的单词学习工具虽然能够提供基本释义，却很难同时解决语境缺失、复习安排随意和学习反馈不足等问题，用户往往在“短时间记住、隔几天又忘”的循环里反复消耗精力。",
    "从工程实践角度看，仅有大模型并不能自动解决记忆效率问题": "不过，单靠接入大模型并不能自然得到一个可用的学习系统。真正影响记忆效率的，除了生成内容是否生动，还包括复习时机是否合适、学习状态是否可记录、错题是否能回看、词汇数据是否可维护。只有把词库、学习进度、测验训练、缓存策略和模型生成链路放进同一套工程结构里，AI 能力才不会停留在“回答一次问题”，而是能够持续服务于“学、记、练、测、评”的完整过程。",
    "在国内研究方面，关于大模型赋能英语学习的讨论主要集中于教育应用潜力": "从国内研究情况看，现有成果更多集中在大模型进入教育场景后的应用可能性，以及语料、提示词、反馈方式对语言学习的影响。相关研究已经说明，生成式模型在例句构造、个性化解释和学习反馈方面具有较强灵活性，但这些讨论大多停留在教学策略、交互模式或实验原型层面，对词汇系统中的词库维护、学习记录、缓存与回填、测验历史和模型治理等工程问题涉及不多。",
    "综合现有研究可以发现，已有成果要么偏重理论探讨": "综合已有文献可以看到，理论层面的研究已经为“大模型辅助语言学习”提供了较充分的背景支持，但面对本科毕设所要求的可运行系统时，仍然存在一个明显空档：如何把模型生成能力稳定地嵌入业务流程，使生成结果能够被缓存、回写、追踪和重复利用。当前项目恰好落在这一工程化问题上，因此本文并不强调抽象层面的“教育变革”，而是重点讨论一个真实仓库如何把大模型、词汇学习和软件系统设计结合起来。",
    "当前项目采用前后端分离架构。前端基于 Vue 3 构建单页应用": "当前项目采用典型的前后端分离方式组织实现。前端以 Vue 3 单页应用为载体，配合 Vue Router 管理页面切换、Pinia 管理登录态和页面状态、Axios 统一封装接口请求；后端以 Spring Boot 3.3 为核心框架，通过 Spring Security 与 JWT 实现身份认证和接口访问控制。这样的分层方式使页面交互、业务逻辑和数据访问能够分别演进，也方便在不改动核心服务的前提下继续补充词库页面、测验页面或管理端能力。",
    "大模型能力是本系统区别于普通词汇学习软件的关键基础": "大模型接入是本项目最具辨识度的一层能力，但代码实现并没有把它当作孤立接口，而是放在了完整的适配链中处理。后端既支持通过 Ollama 调用本地模型，也支持通过 OpenAI 兼容协议接入云端模型；用户可在资料页中切换模型来源和风格偏好，管理员还可以维护全局或私有模型配置。这样做的直接好处是，系统在离线可用性、云端生成质量和后续扩展性之间保留了足够弹性。",
    "系统整体采用浏览器/服务器模式和前后端分离架构": "从整体设计看，系统沿用浏览器/服务器模式展开。浏览器侧负责学习页面与管理页面的交互呈现，Spring Boot 服务负责业务编排与接口输出，MySQL 保存词汇、用户、学习进度和测验记录，Redis 则承担热点缓存、看板缓存和部分状态加速功能；模型适配层独立于普通业务层之外，用于处理本地模型和云端模型的统一调用。这样的结构既保持了学习业务的清晰边界，也让模型接入不会直接侵入词库、复习和测验等核心模块。",
    "用户认证模块以 AuthController 和 AuthService 为入口": "用户认证部分围绕 AuthController、AuthServiceImpl 与前端路由守卫协同工作。后端在登录成功后签发 JWT，前端将令牌写入本地存储并在后续请求中自动携带；当用户访问 `/dashboard`、`/words`、`/review`、`/quiz` 等受保护路由时，前端会先检查登录状态，再根据角色信息决定是否允许进入管理页面。这样一来，普通学习数据和管理员维护能力被清晰分离，系统后续扩展私有模型或批量导入等功能时也更容易控制权限边界。",
    "当前项目测试环境以本地开发环境为主": "本项目的测试与验证工作主要在本地联调环境中完成。后端使用 Spring Boot、Maven、MySQL 和 Redis 组成运行环境，前端使用 Vue 3、Vite 与 Element Plus 进行页面编译和联调。考虑到毕业设计阶段更强调功能闭环与工程真实性，本文采用“自动化测试结果 + 构建校验 + 实际运行截图”的组合方式组织测试证据，而不是仅凭理论性的测试表格描述系统状态。",
    "在自动化测试方面，后端选取了 `SM2AlgorithmTest`": "自动化测试部分以当前仓库中能够稳定复现的测试类为依据，重点覆盖了 SM-2 算法、用户控制器、用户服务、云模型边界和角色权限等模块。其中，SM2AlgorithmTest 直接验证了复习间隔、重复次数和 E-Factor 下限等关键逻辑；CloudLlmModelServiceImplBoundaryTest 则检验了私有模型更新、跨用户访问限制和 API Key 清理等边界情况。配合前端 `npm run build` 的生产构建校验，可以较好证明当前代码至少在核心逻辑和前端工程层面保持可运行状态。",
    "本文以“集成大模型解释能力的四六级高频词记忆系统”为研究对象": "本文围绕当前代码仓库中的 CET46 Vocabulary 项目，对一个“以词汇学习为主线、以大模型解释为增强、以复习调度为核心”的系统进行了梳理与总结。结合任务书、开题报告和现有代码可以看出，项目已经不再只是一个简单的四六级背词页面，而是逐步形成了从词库浏览、词条扩展、复习调度、测验训练到模型配置管理的完整学习闭环。",
    "从工程实现角度看，本项目的特点主要体现在三个方面": "从实现效果来看，项目有三个较为突出的工程特点。其一，AI 能力并未停留在对话式问答，而是被拆解为例句、助记、释义和扩展解释等结构化结果，能够直接服务单词详情页和学习助手界面。其二，系统同时保留本地模型与云端模型两条接入路径，并辅以缓存、异步回填、状态修正和回退逻辑，使模型能力在不稳定环境下仍然具备基本可用性。其三，复习、测验和看板统计与词汇学习数据打通后，用户不只是“看解释”，而是可以持续积累学习轨迹并获得反馈。",
    "当然，当前系统仍然存在进一步完善空间": "同时也应看到，当前系统仍处在适合教学展示和毕业设计总结的阶段，而非面向大规模用户的成熟产品。现有测试更多体现为本地自动化与功能回归，尚缺少更长周期的真实学习效果评估；测验与推荐逻辑也还有继续结合用户画像和错题记录进行细化的空间。后续若能在真实使用数据、提示词优化、模型成本控制和内容审核机制上继续深入，系统的教学辅助价值还可以进一步提升。",
}


ASSET_PLAN = [
    ("4.1 系统总体架构设计", "系统各层之间的组成关系如图4-1所示。", ROOT / "docs/thesis-assets/diagrams/system-architecture.png", "图4-1 系统总体架构图", 15.5),
    ("4.2 功能模块设计", "系统功能模块划分情况如图4-2所示。", ROOT / "docs/thesis-assets/diagrams/functional-modules.png", "图4-2 系统功能模块图", 15.5),
    ("4.3 数据库设计", "数据库核心实体及其关联关系如图4-3所示。", ROOT / "docs/thesis-assets/diagrams/er-diagram.png", "图4-3 系统 E-R 图", 15.5),
    ("5.1 用户认证与权限控制实现", "用户注册后进入学习风格初始化页面，系统运行效果如图5-1所示。", ROOT / "docs/thesis-assets/screenshots/system-onboarding.png", "图5-1 用户注册与学习风格初始化页面", 14.5),
    ("5.2 词库导入与词条管理实现", "词库浏览与加入学习页面的实际运行效果如图5-2所示。", ROOT / "docs/thesis-assets/screenshots/system-word-list.png", "图5-2 词库浏览与加入学习页面", 15.0),
    ("5.3 单词详情页与 AI 生成链路实现", "单词详情页触发 AI 生成、缓存回填与前端轮询展示的关键交互关系如图5-3所示。", ROOT / "docs/thesis-assets/diagrams/ai-generation-sequence.png", "图5-3 单词详情 AI 生成与回填时序图", 15.5),
    ("5.3 单词详情页与 AI 生成链路实现", "单词详情页实际运行效果如图5-4所示。", ROOT / "docs/thesis-assets/screenshots/system-word-detail.png", "图5-4 单词详情与 AI 解释生成页面", 15.0),
    ("5.4 学习助手与本地/云端双模型适配实现", "学习助手页面的运行效果如图5-5所示。", ROOT / "docs/thesis-assets/screenshots/system-assistant.png", "图5-5 学习助手页面运行效果", 15.0),
    ("5.5 复习调度与学习看板实现", "学习看板页面的运行效果如图5-6所示。", ROOT / "docs/thesis-assets/screenshots/system-dashboard.png", "图5-6 学习看板页面运行效果", 15.0),
    ("5.6 模拟测验与历史记录实现", "模拟测验页面的运行效果如图5-7所示。", ROOT / "docs/thesis-assets/screenshots/system-quiz.png", "图5-7 模拟测验页面运行效果", 15.0),
    ("6.2 自动化测试结果分析", "基于测试报告和构建结果整理的证据截图如图6-1所示。", ROOT / "docs/thesis-assets/screenshots/system-test-report.png", "图6-1 自动化测试与构建结果证据截图", 15.0),
    ("6.4 部署过程与运行支撑", "根据当前仓库部署脚本、Nginx 配置和启动日志整理的部署证据截图如图6-2所示。", ROOT / "docs/thesis-assets/screenshots/system-deploy-report.png", "图6-2 部署脚本与启动日志证据截图", 15.0),
    ("6.4 部署过程与运行支撑", "项目部署步骤可概括为如图6-3所示的流程。", ROOT / "docs/thesis-assets/diagrams/deployment-flow.png", "图6-3 系统部署流程图", 14.5),
]


FINAL_REFERENCES = [
    "［1］卢宇, 余京蕾, 陈鹏鹤, 等．多模态大模型的教育应用研究与展望[J]．电化教育研究, 2023, 44(6): 38-44．",
    "［2］王艾艾．语料库数据驱动下的英语四级词汇教学研究[J]．Advances in Education, 2024, 14: 1-10．",
    "［3］张艺凡, 赵静怡, 藕才俊．基于 Ebbinghaus 模型的高效记忆工具设计与实现[J]．计算机科学与应用, 2024, 14(10): 22-32．",
    "［4］陈义嘉, 刘芳．基于大模型技术的个性化外语学习系统[J]．Software Engineering and Applications, 2025, 14: 647-658．",
    "［5］方明炜, 朱君辉, 鲁鹿鸣, 等．例句质量评估体系构建及大语言模型例句生成能力评估[J]．现代教育技术, 2025, 35(2): 88-96．",
    "［6］郭子浩, 孙由之, 张梦林, 等．大语言模型背景下提示词工程赋能英语口语学习研究[J]．Advances in Education, 2023, 13: 8213-8220．",
    "［7］White J, Fu Q, Hays P, et al．A Prompt Pattern Catalog to Enhance Prompt Engineering with ChatGPT[J/OL]．arXiv, 2023．",
    "［8］Wang S, Hu Y, Yang X, et al．Personalized Forgetting Mechanism with Concept-Driven Knowledge Tracing[J/OL]．arXiv, 2024．",
    "［9］Subramani S, Anto J S, Sarensanth C, et al．Learn Buddy: Transforming Education with Generative AI[C]．2024 International Conference on Computing and Data Science, 2024: 1-5．",
    "［10］Fincham N X, Alvarez A A．Using Large Language Models to Facilitate L2 Proficiency Development through Personalized Feedback and Scaffolding[C]．Proceedings of the International CALL Research Conference, 2024: 59-64．",
    "［11］Spring Team．Spring Boot Reference Documentation[EB/OL]．https://docs.spring.io/spring-boot/docs/current/reference/html/，2026-04-06．",
    "［12］Vue Team．Vue Router Documentation[EB/OL]．https://router.vuejs.org/，2026-04-06．",
]


def set_run_fonts(run, size_pt=12, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_paragraph_text(paragraph, text, size_pt=12, bold=False, align=None, first_line_indent=True):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    return new_para


def add_body_paragraph_after(paragraph, text):
    new_para = insert_paragraph_after(paragraph)
    set_paragraph_text(new_para, text, size_pt=12, first_line_indent=True)
    return new_para


def add_picture_after(paragraph, image_path: Path, caption: str, width_cm: float):
    pic_para = insert_paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.first_line_indent = Pt(0)
    pic_para.paragraph_format.line_spacing = Pt(20)
    pic_para.paragraph_format.space_before = Pt(0)
    pic_para.paragraph_format.space_after = Pt(0)
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))

    cap_para = insert_paragraph_after(pic_para)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.first_line_indent = Pt(0)
    cap_para.paragraph_format.line_spacing = Pt(20)
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after = Pt(0)
    run = cap_para.add_run(caption)
    set_run_fonts(run, size_pt=10.5)
    return cap_para


def apply_text_replacements(document):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in TEXT_REPLACEMENTS.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement, size_pt=12, first_line_indent=True)
                break


def update_references(document):
    start = None
    end = None
    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text == "参考文献":
            start = idx + 1
        elif start is not None and text == "致  谢":
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("Reference section not found")

    ref_paragraphs = [p for p in document.paragraphs[start:end] if p.text.strip()]
    if len(ref_paragraphs) < len(FINAL_REFERENCES):
        raise RuntimeError("Not enough reference paragraphs to rewrite")

    for paragraph, ref in zip(ref_paragraphs, FINAL_REFERENCES):
        set_paragraph_text(paragraph, ref, size_pt=10.5, first_line_indent=False)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.hanging_indent = Pt(21)


def insert_assets(document):
    heading_map = {paragraph.text.strip(): paragraph for paragraph in document.paragraphs if paragraph.text.strip()}
    anchors = {}
    for heading, intro, image_path, caption, width_cm in ASSET_PLAN:
        if heading not in heading_map:
            raise RuntimeError(f"Heading not found: {heading}")
        if not image_path.exists():
            raise RuntimeError(f"Image not found: {image_path}")
        anchor = anchors.get(heading, heading_map[heading])
        intro_para = add_body_paragraph_after(anchor, intro)
        caption_para = add_picture_after(intro_para, image_path, caption, width_cm)
        anchors[heading] = caption_para


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    document = Document(str(SOURCE_DOCX))
    apply_text_replacements(document)
    update_references(document)
    insert_assets(document)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
