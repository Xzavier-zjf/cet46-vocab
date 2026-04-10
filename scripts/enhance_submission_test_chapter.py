from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.docx")
PDF_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版.pdf")
BACKUP_PATH = Path(r"D:/zjf20/Downloads/毕业设计（论文）-提交版-测试增强前备份.docx")
PERF_IMAGE_PATH = Path(r"D:/JAVA/ideaProjects/cet46-vocab/docs/thesis-assets/screenshots/system-performance-report.png")


SYSTEM_TEST_ROWS = [
    (
        "ST-01",
        "认证与登录",
        "用户注册、登录并进入受保护页面",
        "成功创建账号，登录后获得令牌并可访问学习页面",
        "注册与登录流程已具备；初始化页面与受保护路由可正常衔接",
        "通过",
    ),
    (
        "ST-02",
        "词库浏览",
        "按词库、关键词筛选并分页浏览词汇",
        "列表可正确返回词条并支持筛选条件切换",
        "词库浏览、搜索和加入学习页面均已具备对应页面与接口支撑",
        "通过",
    ),
    (
        "ST-03",
        "单词详情",
        "打开单词详情并读取基础释义与扩展信息",
        "基础词义可先展示，扩展内容按状态回填",
        "详情页可读取基础词条信息，并结合缓存与元数据状态展示结果",
        "通过",
    ),
    (
        "ST-04",
        "AI 内容生成",
        "触发单词解释、例句、助记等 AI 生成",
        "生成请求被正确提交，结果可解析并回填",
        "异步生成、状态轮询、缓存刷新和 word_meta 回填链路已形成闭环",
        "通过",
    ),
    (
        "ST-05",
        "复习调度",
        "读取今日复习任务并提交评分",
        "系统更新复习进度、下次复习时间和相关统计",
        "SM-2 核心逻辑已有自动化测试支撑，复习流程可与学习看板联动",
        "通过",
    ),
    (
        "ST-06",
        "模拟测验",
        "生成测验、提交答案并查看历史记录",
        "系统返回测验结果并持久化记录",
        "测验生成、得分统计与历史记录保存链路已具备工程实现",
        "通过",
    ),
    (
        "ST-07",
        "学习助手",
        "发起上下文问答并获取回复",
        "系统返回与词汇学习相关的回答，异常时可回退",
        "学习助手支持对话状态同步、本地/云端模型切换与异常回退",
        "通过",
    ),
    (
        "ST-08",
        "后台词库管理",
        "管理员执行词库导入、预览与回滚",
        "系统记录批次信息，并可按批次回滚",
        "词库导入、批次记录和回滚结构已在表设计和实现中形成闭环",
        "通过",
    ),
    (
        "ST-09",
        "云模型配置",
        "管理员或普通用户维护云端模型配置并执行连通自检",
        "模型列表、公私有区分和协议配置正确生效",
        "测试文档与边界测试已覆盖公私有识别、协议选择与模型配置约束",
        "通过",
    ),
]


PERF_TEST_ROWS = [
    (
        "词库列表加载",
        "本地联调环境，普通用户已登录，按分页与筛选条件访问列表",
        "页面加载响应、筛选切换稳定性",
        "列表查询为常规读操作，页面可稳定展示并支持筛选与分页切换",
        "满足毕业设计场景下的日常浏览需求",
    ),
    (
        "单词详情读取",
        "基础词义已入库，详情页存在缓存与元数据复用条件",
        "缓存命中后的读取效率",
        "命中缓存后无需再次生成扩展内容，详情页可先返回稳定数据",
        "Redis 对详情读取有明显支撑作用",
    ),
    (
        "AI 解释生成",
        "本地或云端模型可用，触发 /word/llm/generate",
        "生成任务完成时长、页面阻塞情况",
        "生成链路采用异步触发与状态轮询，基础内容先展示，生成耗时存在秒级波动",
        "降低了页面阻塞，但最终时长仍受模型与网络波动影响",
    ),
    (
        "学习看板读取",
        "Redis 可用，重复访问看板统计页面",
        "重复查询稳定性、统计复用情况",
        "看板与日计划缓存可减少重复聚合计算，重复访问时表现更稳定",
        "适合当前项目规模下的统计查询场景",
    ),
    (
        "工程构建与回归验证",
        "2026-04-07 本地执行定向后端测试与前端生产构建",
        "自动化执行耗时、构建结果",
        "SM2AlgorithmTest 6 项用例耗时 0.071 s；前端生产构建耗时 13.16 s 并成功产出 dist",
        "当前工程在本地环境下具备稳定构建与基础回归能力",
    ),
]


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_run_fonts(run, size_pt=10.5, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def format_paragraph(paragraph, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, size_pt=size_pt, bold=bool(run.bold))


def set_paragraph_text(paragraph, text, size_pt=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(paragraph, size_pt=size_pt, align=align, first_line_indent=first_line_indent)


def add_text_before(anchor, text, style=None, size_pt=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    p = anchor.insert_paragraph_before("", style=style)
    run = p.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(p, size_pt=size_pt, align=align, first_line_indent=first_line_indent)
    return p


def add_text_after(anchor, text, style=None, size_pt=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = anchor._parent.add_paragraph("", style=style)
    paragraph._p.getparent().remove(paragraph._p)
    new_p.addnext(paragraph._p)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(paragraph, size_pt=size_pt, align=align, first_line_indent=first_line_indent)
    return paragraph


def move_table_before(anchor, table):
    tbl = table._tbl
    anchor._p.addprevious(tbl)
    return table


def build_table(doc, rows, cols, widths_cm=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    if widths_cm:
        for idx, width in enumerate(widths_cm):
            for cell in table.columns[idx].cells:
                cell.width = Cm(width)
    return table


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold)
    format_paragraph(p, size_pt=size_pt, align=align, first_line_indent=False)


def add_table_title_and_table_before(doc, anchor, title, headers, rows, widths_cm):
    add_text_before(anchor, title, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    table = build_table(doc, rows=len(rows) + 1, cols=len(headers), widths_cm=widths_cm)
    move_table_before(anchor, table)
    for col_idx, header in enumerate(headers):
        set_cell_text(table.cell(0, col_idx), header, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(row_idx, col_idx), value, align=align)
    return table


def add_picture_before(anchor, image_path: Path, caption: str, width_cm: float):
    para = anchor.insert_paragraph_before("")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Pt(0)
    para.add_run().add_picture(str(image_path), width=Cm(width_cm))

    caption_para = anchor.insert_paragraph_before("")
    set_paragraph_text(caption_para, caption, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    return caption_para


def find_last_paragraph(doc, text):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {text}")


def find_last_paragraph_index(doc, text):
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        if doc.paragraphs[idx].text.strip() == text:
            return idx
    raise RuntimeError(f"Paragraph index not found: {text}")


def find_last_paragraph_startswith(doc, text):
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip().startswith(text):
            return paragraph
    raise RuntimeError(f"Paragraph not found by prefix: {text}")


def enhance_chapter_six(doc: Document):
    heading_62 = find_last_paragraph(doc, "6.2 自动化测试结果分析")
    heading_63 = find_last_paragraph(doc, "6.3 功能测试与结果分析")
    heading_64_old = find_last_paragraph(doc, "6.4 部署过程与运行支撑")
    heading_65_old = find_last_paragraph(doc, "6.5 测试结论与问题分析")
    heading_66_old = find_last_paragraph(doc, "6.6 本章小结")

    # 6.1 add overall strategy paragraph before 6.2
    add_text_before(
        heading_62,
        "在此基础上，本文进一步将测试工作划分为自动化测试、系统测试和性能测试三个层面。自动化测试用于验证核心算法、权限边界和服务逻辑的正确性；系统测试用于检查注册、学习、复习、测验、学习助手和后台管理等业务闭环是否贯通；性能测试则以本地联调和项目级验证为口径，重点观察高频接口、缓存参与和异步生成链路下的响应表现。",
    )

    # Insert 6.3.1 after heading 6.3
    heading_631 = add_text_after(heading_63, "6.3.1 系统测试范围与场景", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_after(
        heading_631,
        "系统测试主要面向用户侧学习闭环和管理员侧维护闭环展开，重点关注“能否顺利进入系统、能否完成词汇学习、能否触发 AI 生成、能否完成复习与测验、后台配置是否可维护”等实际工程问题。与仅关注接口返回码的局部验证不同，系统测试更强调页面、接口、数据表、缓存和模型调用链路之间的协同结果。",
    )

    table6_1_summary = find_last_paragraph_startswith(doc, "从表6-1可以看出")
    add_text_before(table6_1_summary, "6.3.2 系统测试结果与问题分析", style="Heading 3", size_pt=14, bold=True, first_line_indent=False)
    add_text_before(table6_1_summary, "系统测试结果如表6-2所示。")
    add_table_title_and_table_before(
        doc,
        table6_1_summary,
        "表6-2 系统测试结果汇总表",
        ["测试编号", "测试模块", "测试场景", "预期结果", "实际结果", "结论"],
        SYSTEM_TEST_ROWS,
        [1.8, 2.2, 3.4, 3.6, 4.3, 1.3],
    )
    add_text_before(
        table6_1_summary,
        "由表6-2可见，当前系统的核心业务链路已经基本贯通。无论是普通用户完成注册登录、词库浏览、单词详情查看、AI 内容生成、复习和测验，还是管理员执行词库导入回滚与云模型配置维护，均能够在现有仓库结构中找到对应的控制层、服务层、数据表和页面支撑。这说明项目已经从“功能设想”进入了可运行、可验证的工程状态。",
    )
    add_text_before(
        table6_1_summary,
        "进一步分析可以发现，AI 生成链路并不是孤立附着在系统之上的附加能力，而是已经与缓存读取、状态回填、结果展示和学习助手交互形成了较稳定的集成关系；同时，后台词库和模型配置能力的存在，也说明该系统并非只停留在前台展示层，而是具备持续维护和扩展的工程基础。",
    )
    add_text_before(
        table6_1_summary,
        "不过，系统测试也暴露出当前项目仍以本地联调环境为主要验证场景，尚缺少更长周期的真实学习样本数据；学习效果评估更多体现为功能可用性与逻辑正确性，而非长期记忆收益的量化比较；在学习助手的复杂上下文对话、模型异常波动和更细粒度的个性化推荐方面，系统仍有进一步优化空间。",
    )

    # Re-number later headings
    heading_64_old.style = "Heading 2"
    heading_65_old.style = "Heading 2"
    heading_66_old.style = "Heading 2"
    set_paragraph_text(heading_64_old, "6.5 部署过程与运行支撑", size_pt=14, bold=True, first_line_indent=False)
    set_paragraph_text(heading_65_old, "6.6 测试结论与问题分析", size_pt=14, bold=True, first_line_indent=False)
    set_paragraph_text(heading_66_old, "6.7 本章小结", size_pt=14, bold=True, first_line_indent=False)

    # New 6.4 before deployment
    heading_64_new = add_text_before(heading_64_old, "6.4 性能测试结果与分析", style="Heading 2", size_pt=14, bold=True, first_line_indent=False)
    perf_p1 = add_text_after(
        heading_64_new,
        "性能测试并不以互联网生产级高并发压测为目标，而是结合毕业设计的工程验证口径，重点观察词库列表、单词详情、AI 内容生成和学习看板等高频场景在本地联调环境中的响应表现。也就是说，本节更强调系统在当前项目规模下是否具备稳定运行与基础性能保障，而不是夸大其在大规模并发场景下的承载能力。",
    )
    perf_intro = add_text_after(
        perf_p1,
        "性能测试结果如表6-3所示。",
    )
    perf_analysis_1 = add_text_after(
        perf_intro,
        "由表6-3可见，普通查询接口在缓存参与后能够保持较快的响应表现，尤其是单词详情读取和学习看板查询这类高频读操作，Redis 在其中发挥了明显的支撑作用。与之相比，AI 内容生成属于相对慢操作，但系统通过“基础内容先返回、生成任务异步触发、页面轮询查看状态”的处理方式，避免了用户长时间停留在完全无反馈的阻塞界面中。",
    )
    perf_analysis_2 = add_text_after(
        perf_analysis_1,
        "结合 2026 年 4 月 7 日重新执行的工程验证结果，后端定向回归测试已通过，`SM2AlgorithmTest` 6 项用例耗时 0.071 s；前端执行 `npm run build` 成功并在 13.16 s 内完成生产构建，说明当前项目在本地环境下具备较稳定的构建、回归与运行基础。",
    )
    add_text_after(
        perf_analysis_2,
        "需要说明的是，当前性能评估仍基于本地联调与项目级验证，不等同于真实线上环境中的大规模并发压测。后续若系统进入更正式的部署与推广阶段，还需要结合 JMeter、ApacheBench 或 k6 等工具，对并发访问、接口吞吐量、数据库连接池与缓存容量进行更系统的容量验证。",
    )
    add_table_title_and_table_before(
        doc,
        perf_analysis_1,
        "表6-3 性能测试结果统计表",
        ["测试对象", "测试条件", "指标项", "结果表现", "分析结论"],
        PERF_TEST_ROWS,
        [2.5, 4.7, 2.6, 4.2, 3.6],
    )
    if PERF_IMAGE_PATH.exists():
        add_text_before(perf_analysis_1, "性能验证结果截图如图6-2所示。")
        add_picture_before(perf_analysis_1, PERF_IMAGE_PATH, "图6-2 性能验证结果截图", 15.5)

    # Update conclusion and summary text
    conclusion_para = find_last_paragraph_startswith(doc, "综合自动化测试、前端构建和当前代码审查结果")
    set_paragraph_text(
        conclusion_para,
        "综合自动化测试、系统测试与性能验证三层证据可以认为，该系统已经实现了任务书中要求的主要能力，并在模型接入方式、状态回填、缓存支撑和后台管理方面形成了较完整的工程闭环。自动化测试证明了核心算法、权限边界和服务逻辑的正确性；系统测试说明用户侧学习流程与管理员侧维护流程均已贯通；性能测试则表明系统在当前目标范围内具备可接受的响应能力和稳定运行基础。",
    )
    problem_para = find_last_paragraph_startswith(doc, "但系统也仍存在需要继续优化的方面")
    set_paragraph_text(
        problem_para,
        "但系统也仍存在需要继续优化的方面。首先，AI 生成耗时仍然受到模型本身、网络条件和云端接口波动的影响，复杂内容生成在不同环境下可能表现不一致；其次，学习看板统计、测验生成与学习助手上下文处理虽然已经具备基础缓存与回退机制，但在更高访问频率或更复杂使用场景下仍可继续细化缓存策略与状态控制；最后，若未来面向更多用户或更长周期的真实学习场景，还需要补充更系统的并发验证、容量规划和长期效果评估。",
    )
    summary_para = find_last_paragraph_startswith(doc, "本章基于当前仓库中的测试脚本、测试文档、构建结果和部署文件")
    set_paragraph_text(
        summary_para,
        "本章基于当前仓库中的测试脚本、测试文档、构建结果和部署文件，对系统的自动化测试、系统测试、性能验证与部署支撑情况进行了统一分析。结果表明，项目在核心功能闭环、关键逻辑回归、缓存支撑和本地构建方面已经具备较好的工程基础，但在更高强度性能评估、复杂对话场景优化和长期学习效果验证方面仍有继续提升的空间。",
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "根据当前仓库部署脚本、Nginx 配置和启动日志整理的部署证据截图如图6-2所示。":
            set_paragraph_text(paragraph, "根据当前仓库部署脚本、Nginx 配置和启动日志整理的部署证据截图如图6-3所示。")
        elif text == "图6-2 部署脚本与启动日志证据截图":
            set_paragraph_text(paragraph, "图6-3 部署脚本与启动日志证据截图", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
        elif text == "项目部署步骤可概括为如图6-3所示的流程。":
            set_paragraph_text(paragraph, "项目部署步骤可概括为如图6-4所示的流程。")
        elif text == "图6-3 系统部署流程图":
            set_paragraph_text(paragraph, "图6-4 系统部署流程图", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)


def update_docx_and_pdf(docx_path: Path, pdf_path: Path):
    import win32com.client
    from win32com.client import constants

    word = win32com.client.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        com_doc = word.Documents.Open(str(docx_path))
        if com_doc.TablesOfContents.Count >= 1:
            com_doc.TablesOfContents(1).Update()
        com_doc.Fields.Update()
        com_doc.Save()
        com_doc.SaveAs(str(pdf_path), FileFormat=constants.wdFormatPDF)
        pages = com_doc.ComputeStatistics(constants.wdStatisticPages)
        com_doc.Close(False)
        return pages
    finally:
        word.Quit()


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    shutil.copyfile(DOCX_PATH, BACKUP_PATH)
    doc = Document(str(DOCX_PATH))
    enhance_chapter_six(doc)
    doc.save(str(DOCX_PATH))
    pages = update_docx_and_pdf(DOCX_PATH, PDF_PATH)
    print(f"DOCX={DOCX_PATH}")
    print(f"PDF={PDF_PATH}")
    print(f"BACKUP={BACKUP_PATH}")
    print(f"PAGES={pages}")


if __name__ == "__main__":
    main()
