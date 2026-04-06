from copy import deepcopy
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_DOCX = r"C:/Users/zjf20/Desktop/毕业设计（论文）.docx"
OUTPUT_DOCX = r"C:/Users/zjf20/Desktop/毕业设计（论文）-项目重写稿.docx"
TITLE = "集成大模型解释能力的四六级高频词记忆系统设计与实现"


def clear_paragraph(paragraph):
    if hasattr(paragraph, "clear"):
        paragraph.clear()
        return
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_run_fonts(run, size_pt, bold=False, font_cn="宋体", font_en="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font_cn)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)


def set_paragraph_text(
    paragraph,
    text,
    size_pt=12,
    bold=False,
    align=None,
    first_line_indent=True,
    font_cn="宋体",
    font_en="Times New Roman",
):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, size_pt=size_pt, bold=bold, font_cn=font_cn, font_en=font_en)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line_indent else Pt(0)
    return paragraph


def add_paragraph(
    document,
    text,
    size_pt=12,
    bold=False,
    align=None,
    first_line_indent=True,
    font_cn="宋体",
    font_en="Times New Roman",
):
    paragraph = document.add_paragraph()
    return set_paragraph_text(
        paragraph,
        text,
        size_pt=size_pt,
        bold=bold,
        align=align,
        first_line_indent=first_line_indent,
        font_cn=font_cn,
        font_en=font_en,
    )


def add_heading1(document, text):
    p = document.add_paragraph(style="Heading 1")
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_fonts(run, size_pt=16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_heading2(document, text):
    p = document.add_paragraph(style="Heading 2")
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_fonts(run, size_pt=15, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_heading3(document, text):
    p = document.add_paragraph(style="Heading 3")
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_fonts(run, size_pt=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    return p


def add_page_break(document):
    p = document.add_paragraph()
    p.add_run().add_break()


def add_toc_field(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = Pt(0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text_run = OxmlElement("w:r")
    text_t = OxmlElement("w:t")
    text_t.text = "右键更新域后生成三级目录"
    text_run.append(text_t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r1 = p.add_run()
    r1._r.append(fld_begin)
    r2 = p.add_run()
    r2._r.append(instr)
    r3 = p.add_run()
    r3._r.append(fld_sep)
    r4 = p.add_run()
    r4._r.append(text_run)
    r5 = p.add_run()
    r5._r.append(fld_end)
    for run in p.runs:
        set_run_fonts(run, size_pt=12, bold=False)


def add_references(document, refs):
    add_heading1(document, "参考文献")
    for ref in refs:
        p = add_paragraph(document, ref, size_pt=10.5, first_line_indent=False)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.hanging_indent = Pt(21)


def add_ack(document):
    add_page_break(document)
    add_heading1(document, "致  谢")
    add_paragraph(
        document,
        "毕业设计完成过程中，指导教师在选题论证、系统实现思路、论文结构组织和阶段性修改方面给予了持续指导与帮助，本文在此表示诚挚感谢。由于学校材料要求涉及个人信息和签名内容，相关姓名与签署信息在最终提交前由本人按学校规定补充完善。",
        size_pt=10.5,
    )
    add_paragraph(
        document,
        "同时，感谢学院提供的学习环境与项目实践条件，感谢课程学习过程中各位任课教师所奠定的软件工程、数据库、Web 开发与人工智能应用基础。项目开发和论文撰写过程中，仓库中的现有代码、测试记录、任务书与开题报告共同构成了本文的重要依据，使系统分析、设计、实现和测试内容能够建立在真实工程证据之上。",
        size_pt=10.5,
    )
    add_paragraph(
        document,
        "最后，感谢在资料整理、环境搭建和功能联调过程中给予支持与建议的同学和朋友。后续本人将继续根据导师意见和答辩反馈，对图表细节、参考文献时效性以及个别章节篇幅进行进一步打磨，以形成更加完整、规范的毕业设计定稿。",
        size_pt=10.5,
    )


def add_content(document):
    add_page_break(document)
    add_heading1(document, "摘  要")
    add_paragraph(
        document,
        "本文围绕“集成大模型解释能力的四六级高频词记忆系统”展开设计与实现，目标在于构建一个兼顾词汇数据管理、智能解释生成、复习调度、模拟测验和学习统计的英语词汇学习平台。系统采用前后端分离架构，前端基于 Vue3、Vite、Element Plus 与 Pinia 实现学习页、词库页、复习页、测验页、学习助手页和管理页等界面；后端基于 Spring Boot 3.3、Spring Security、MyBatis-Plus、MySQL、Redis 和 JWT 完成认证授权、词汇管理、AI 内容生成、复习调度与测验记录等核心服务。与任务书中“以本地轻量化模型为主”的设想相比，当前项目进一步实现了本地 Ollama 与云端 OpenAI 兼容接口的双模型接入能力，并通过异步生成、缓存回填、JSON 结构化解析、模型可用性检测和私有模型配置管理，提升了学习内容生成的稳定性与可维护性。系统在功能上实现了 CET-4/CET-6 词库浏览、单词详情展示、智能例句与助记生成、SM-2 间隔重复复习、选择题/填空题测验、学习数据看板、词库导入及模型使用统计等模块。结合现有自动化测试、前端构建校验和手工测试用例整理结果可知，系统核心流程能够正常运行，具备较好的工程完整性。本文重点介绍了系统需求分析、总体设计、数据库设计、关键模块实现、测试与部署过程，并对系统的优势、局限及后续优化方向进行了总结。",
    )
    add_paragraph(
        document,
        "关键词：四六级词汇；大语言模型；SM-2算法；智能学习系统；前后端分离",
        first_line_indent=False,
    )

    add_page_break(document)
    add_heading1(document, "Abstract")
    add_paragraph(
        document,
        "This paper presents the design and implementation of a CET-4 and CET-6 high-frequency vocabulary memorization system integrated with large language model based explanation capability. The system aims to provide an engineering-oriented learning platform that combines vocabulary data management, AI generated explanation, review scheduling, quiz training, and learning analytics. The front end is built with Vue 3, Vite, Element Plus and Pinia, while the back end is implemented with Spring Boot 3.3, Spring Security, MyBatis-Plus, MySQL, Redis and JWT. Compared with the original proposal that mainly emphasized local lightweight models, the current project further supports both local Ollama models and cloud models with OpenAI compatible protocols. By introducing asynchronous generation, cache backfilling, structured JSON parsing, model availability checking and private cloud model management, the system improves the reliability of generated learning content. The implemented functions include CET-4 and CET-6 vocabulary browsing, word detail display, AI generated example sentences and mnemonics, SM-2 based spaced repetition review, multiple choice and fill in the blank quizzes, dashboard analytics, vocabulary import, and model usage management. Existing automated tests, front-end production build verification, and documented manual test cases indicate that the major workflows of the system are executable and maintainable. This thesis focuses on requirement analysis, overall architecture, database design, implementation of key modules, testing, deployment process, and future optimization directions.",
        font_cn="Times New Roman",
        font_en="Times New Roman",
    )
    add_paragraph(
        document,
        "Key words: CET vocabulary; large language model; SM-2 algorithm; intelligent learning system; front-end and back-end separation",
        first_line_indent=False,
    )

    add_page_break(document)
    add_heading1(document, "目  录")
    add_toc_field(document)

    add_page_break(document)
    add_heading1(document, "第一章 绪论")
    add_heading2(document, "1.1 选题背景与研究意义")
    add_paragraph(
        document,
        "大学英语四、六级考试仍然是高校英语学习中的重要评价节点，词汇掌握效果直接影响阅读理解、听力理解、写作和翻译等多类题型表现。传统的单词书和静态词典类应用虽然能够提供基础释义，却普遍存在学习过程割裂、语境支持不足和复习节奏不科学等问题，导致学习者在“记得快、忘得也快”的状态下反复投入大量时间成本。",
    )
    add_paragraph(
        document,
        "近年来，生成式人工智能和大语言模型在教育场景中的应用价值逐步显现。相关研究指出，大模型可以在例句生成、个性化解释、学习反馈和内容重构等方面提供超越固定资源库的灵活支持[1][4]。同时，提示工程方法为控制模型输出结构和提升结果可用性提供了新的路径，使模型能够更好地服务于可编程系统而非停留在纯聊天式交互层面[6][7]。",
    )
    add_paragraph(
        document,
        "从工程实践角度看，仅有大模型并不能自动解决记忆效率问题，仍需要将学习内容生成能力与科学的复习调度机制结合起来。基于艾宾浩斯遗忘规律演化而来的间隔重复方法已在单词记忆软件中得到验证[3][5]，如果能将其与结构化词库、学习日志、测验记录和智能解释服务协同设计，便可以形成覆盖“学、记、练、测、评”的完整学习闭环。",
    )
    add_paragraph(
        document,
        "因此，设计并实现一个集成大模型解释能力的四六级高频词记忆系统，不仅具有现实应用意义，也具备较强的软件工程研究价值。一方面，该系统能够面向真实学习场景提升词汇理解和复习效率；另一方面，它为本地模型接入、云端模型兼容、缓存策略、JSON 解析、复习算法落地和前后端协同提供了完整工程案例。",
    )

    add_heading2(document, "1.2 国内外研究现状")
    add_paragraph(
        document,
        "在国内研究方面，关于大模型赋能英语学习的讨论主要集中于教育应用潜力、提示词设计和个性化教学策略。卢宇等从教育技术视角分析了多模态大模型对教育内容生产方式的影响，指出其能够推动学习资源从统一供给走向个性化生成[1]；王艾艾从英语词汇教学角度说明语料驱动与数字化学习资源对于提升词汇教学效果具有积极意义[2]；郭子浩等则从提示词工程角度讨论了大模型在语言学习中的组织方式和交互方式[6]。",
    )
    add_paragraph(
        document,
        "在国外研究方面，学者更关注生成式 AI 在学习支持、反馈增强和结构化输出控制上的可操作性。White 等提出的 Prompt Pattern Catalog 系统总结了角色模式、模板模式等常见提示模式，为开发者将大模型嵌入软件系统提供了较强的方法参考[7]；Subramani 等在教育软件场景下讨论了生成式 AI 对学习过程的重构价值[9]；Fincham 与 Alvarez 进一步从二语学习角度分析了大语言模型在个性化反馈中的应用潜力[10]。",
    )
    add_paragraph(
        document,
        "综合现有研究可以发现，已有成果要么偏重理论探讨，要么集中在聊天式问答原型，对“结构化词汇学习系统”这一具体场景中的词库管理、复习调度、测验训练、数据统计和模型配置管理关注不足。特别是在本科毕设可落地的工程尺度下，如何把模型生成内容转化为稳定、可缓存、可追踪、可重复调用的学习资源，仍然是一个值得深入实现和总结的问题。",
    )

    add_heading2(document, "1.3 本文研究内容与目标")
    add_paragraph(
        document,
        "结合任务书、开题报告和当前仓库实现，本文研究目标可归纳为四个方面。第一，建立覆盖 CET-4 与 CET-6 高频词汇的标准化词库结构，并支持管理员通过批量导入方式维护词汇数据和导入批次记录。第二，构建面向单词详情页和学习助手的 AI 解释生成链路，使系统能够生成例句、近义词辨析、助记信息和扩展解释，并支持本地模型与云端模型两种接入方式。",
    )
    add_paragraph(
        document,
        "第三，将 SM-2 间隔重复算法嵌入学习进度管理过程，根据用户在复习环节中的评分动态更新复习间隔、重复次数和下一次复习日期，进而形成“加入学习—进入复习队列—记录学习状态—在看板中反馈结果”的闭环。第四，在学习功能之外补充模拟测验、学习看板、模型使用统计和模型配置管理等配套模块，使系统不仅能生成学习内容，还能支撑阶段性训练和学习行为分析。",
    )
    add_paragraph(
        document,
        "与开题报告相比，当前项目代码的实际落地范围更加偏向系统工程实现，尤其体现在云端模型兼容、私有模型配置、异步任务重试、使用量统计和用户资料中的模型偏好管理等方面。这些内容在任务书中没有展开，但均来自当前仓库的真实模块，因此本文在章节安排上将以项目代码为最高优先级，任务书和开题报告作为边界约束和目标对照。",
    )

    add_heading2(document, "1.4 论文结构")
    add_paragraph(
        document,
        "第一章为绪论，介绍研究背景、研究意义、国内外研究现状、本文的研究内容与章节安排。第二章介绍系统实现所依赖的相关技术与理论基础，包括前后端框架、大模型接入方式、提示结构化输出和 SM-2 算法。第三章结合当前项目的实际功能进行系统需求分析，明确角色、业务流程和功能边界。第四章给出系统总体设计方案，包括架构分层、功能模块设计、数据库设计和接口设计。第五章围绕关键模块展开详细设计与实现分析，重点说明认证授权、词库导入、AI 生成、学习助手、复习调度和测验流程。第六章介绍系统测试与部署过程，给出当前仓库可支撑的测试结果与分析。最后在结论部分总结全文工作并讨论后续优化方向。",
    )
    add_heading2(document, "1.5 本章小结")
    add_paragraph(
        document,
        "本章从英语词汇学习的现实痛点、大模型教育应用趋势和软件工程落地需求三个层面阐述了本课题的研究背景与意义，梳理了与本课题相关的国内外研究现状，并在任务书与开题报告基础上结合当前代码实现明确了本文的研究目标与章节结构。上述内容为后续的技术分析、需求建模和系统设计提供了问题背景与研究依据。",
    )

    add_page_break(document)
    add_heading1(document, "第二章 相关技术与理论基础")
    add_heading2(document, "2.1 前后端分离开发架构")
    add_paragraph(
        document,
        "当前项目采用前后端分离架构。前端基于 Vue 3 构建单页应用，通过 Vue Router 负责页面导航与访问守卫，通过 Pinia 保存用户身份、主题状态和学习会话信息，通过 Axios 统一封装请求和身份令牌；后端基于 Spring Boot 3.3 提供 RESTful 服务接口，通过 Spring Security 与 JWT 实现登录认证和接口授权控制。该架构能够将交互层与业务层职责解耦，便于在不影响后端逻辑的前提下独立优化前端页面体验[11][12]。",
    )
    add_paragraph(
        document,
        "从当前代码结构看，前端已形成 Dashboard、Learn、WordList、WordDetail、Review、Quiz、Assistant、Statistics、Profile 以及管理端页面等多个视图模块，后端则围绕 Auth、Word、Review、Quiz、Assistant、Dashboard、Admin 等控制器展开组织。页面访问统一通过路由守卫检查登录状态、角色信息和模型偏好初始化状态，后端服务则按 Controller、Service、Mapper 分层组织，具备较好的模块边界与扩展性。",
    )

    add_heading2(document, "2.2 大模型接入与结构化输出控制")
    add_paragraph(
        document,
        "大模型能力是本系统区别于普通词汇学习软件的关键基础。后端同时实现了本地模型与云端模型的接入路径：本地模型通过 Ollama 客户端完成文本生成，云端模型通过 OpenAI 兼容协议完成接口调用；用户可在资料页中选择模型提供方、模型名称和风格偏好，管理员则可维护全局模型和私有模型配置。这种“双通道”方案既满足离线部署和隐私保护需求，也为高质量生成提供了更灵活的替代路径。",
    )
    add_paragraph(
        document,
        "为了避免大模型输出过于随意而难以直接用于程序展示，系统在单词详情和学习助手模块中都引入了结构化提示设计。项目中的 PromptTemplate、LlmResponseParser 和相关 DTO 将模型输出约束在明确字段范围内，使结果能够落入例句、近义词、助记信息和扩展解释等结构中。这种做法与提示模式理论中“模板模式”和“角色模式”的思想一致，能够降低模型幻觉对前端展示和后端落库的影响[7]。",
    )
    add_paragraph(
        document,
        "此外，代码中还实现了缓存回填、生成状态标记、卡住任务修复和待处理任务重试机制。例如 WordServiceImpl 会根据 genStatus、aiExplainStatus 和更新时间判断是否需要再次触发异步生成，既避免页面长时间读取空数据，又保证在模型暂时不可用时系统仍能返回兜底内容。相比开题阶段单纯强调“调用模型生成”的设想，当前实现更接近可上线系统的稳态设计。",
    )

    add_heading2(document, "2.3 SM-2 间隔重复算法")
    add_paragraph(
        document,
        "SM-2 算法是复习调度模块的理论核心。其基本思想是根据学习者对单次复习结果的反馈动态调整难度因子、复习间隔和重复次数，使掌握程度较高的词汇以更长间隔进入下一轮复习，而掌握较差的词汇则尽快回到近期复习队列。该方法能够在有限时间下提高复习效率，是数字化记忆系统中应用较为广泛的调度策略[5][8]。",
    )
    add_paragraph(
        document,
        "在当前项目中，SM2Algorithm 类对 score、easiness、interval 和 repetition 的更新进行了实现：当评分低于 3 时视为遗忘，间隔重置为 1 天且重复次数归零；当评分较高时，系统基于难度因子更新下一次复习时间，第一次和第二次成功复习分别采用 1 天和 6 天的固定间隔，此后使用“当前间隔乘以新的 E-Factor”进行推算。这一实现与经典 SM-2 规则保持一致，同时便于通过测试用例进行验证。",
    )
    add_paragraph(
        document,
        "与传统固定周期复习不同，SM-2 的优势在于能够根据学习状态自适应调整复习节奏。项目中的 ReviewServiceImpl 在用户复习提交后会调用算法结果更新 user_word_progress 表，再同步写入 review_log 和看板缓存数据，使算法执行结果能够直接体现在待复习数量、连续学习天数和学习统计面板上，形成从算法到交互展示的完整闭环。",
    )

    add_heading2(document, "2.4 缓存、统计与部署支撑技术")
    add_paragraph(
        document,
        "系统在性能和可维护性方面还依赖 MySQL、Redis、Maven、Vite 和 Nginx 等基础设施。MySQL 用于保存词库数据、用户信息、学习进度、测验历史和模型配置；Redis 主要用于缓存单词详情、看板概览和异步生成结果，降低高频查询时的数据库压力。DashboardController 中还引入了 daily_plan_cache，以减少每日看板数据重复计算的开销。",
    )
    add_paragraph(
        document,
        "部署方面，仓库根目录提供了 deploy.sh 和 nginx.conf 两份部署参考文件。deploy.sh 展示了“拉取代码、前端构建、后端打包、停止旧进程、启动新 Jar、重载 Nginx”的完整流程，nginx.conf 则配置了前端 history 路由回退和 /api 反向代理规则。这些内容说明当前项目不仅完成了开发态实现，也具备较清晰的部署路径。",
    )
    add_heading2(document, "2.5 本章小结")
    add_paragraph(
        document,
        "本章围绕系统实现所依赖的核心技术与理论基础展开说明，分别介绍了前后端分离架构、大模型接入与结构化输出控制、SM-2 间隔重复算法以及缓存和部署支撑技术。通过这些技术的协同，系统能够同时满足交互体验、内容生成、学习调度、数据持久化与工程部署等方面的要求，为后续需求分析和总体设计提供了技术前提。",
    )

    add_page_break(document)
    add_heading1(document, "第三章 系统需求分析")
    add_heading2(document, "3.1 系统目标与用户角色")
    add_paragraph(
        document,
        "结合当前代码与任务书要求，系统总体目标可以概括为“为四六级备考用户提供以词汇为核心、以智能解释为增强、以复习调度为主线的学习平台”。其目标并非取代完整英语教学平台，而是在词汇学习这一高频、可量化、可反复训练的环节上实现更高的学习效率和更好的工程可控性。",
    )
    add_paragraph(
        document,
        "从角色划分看，系统包含普通学习用户和管理员两类核心角色。普通用户负责注册登录、选择词库、查看单词详情、加入学习、进行复习、发起测验、查看统计和使用学习助手；管理员负责用户管理、词库导入、AI 生成复核、云端模型维护和权限配置。当前前端路由中 `/admin`、`/admin/users` 和 `/admin/permissions` 等路径均显式限定为管理员角色，体现了较清晰的权限边界。",
    )

    add_heading2(document, "3.2 功能需求分析")
    add_heading3(document, "3.2.1 词库与词条管理需求")
    add_paragraph(
        document,
        "系统需要支持 CET-4、CET-6 词汇的分页浏览、关键词检索、词性筛选和详情查看。词条详情不仅要展示英文、音标和中文释义，还要根据当前用户偏好的模型风格动态展示例句、近义词、助记信息、扩展解释和语法提示。对于尚未学习的单词，用户还应能够一键加入学习计划，以便后续进入复习队列。",
    )
    add_paragraph(
        document,
        "管理端还需要支持词库批量导入、导入预览、批次记录查看与回滚。AdminWordBankController 中已经实现了 CSV 文件预览、导入、批次管理和回滚逻辑，同时将导入变化同步记录在 word_import_batch 和 word_import_batch_item 表中。这说明当前仓库的实际实现比开题报告中“构建词库并导入”更进一步，已经覆盖了数据维护审计能力。",
    )
    add_heading3(document, "3.2.2 智能学习内容生成需求")
    add_paragraph(
        document,
        "系统的智能学习内容生成需求主要体现在两个方面：其一，单词详情页需要围绕指定词条生成与学习直接相关的内容，包括英文例句、中文释义、近义词辨析、助记方式和拓展解释；其二，学习助手需要结合当前词条上下文和最近对话历史，回答用户关于词义、搭配、记忆方法和备考策略的问题。为保证系统可用性，生成结果必须可以被前端结构化展示，而不是停留在无边界的聊天文本。",
    )
    add_paragraph(
        document,
        "同时，模型配置需要具备灵活性。根据当前项目代码，用户可以选择本地模型或云端模型，管理员可以维护全局模型列表，系统还支持私有云端模型、连通性自检和协议类型配置。这些需求在任务书中并未完全展开，但它们直接影响生成模块的实际可用性，因此应纳入系统功能需求范围。",
    )
    add_heading3(document, "3.2.3 复习、测验与统计需求")
    add_paragraph(
        document,
        "复习模块需要根据用户的学习进度生成当日待复习队列，并在复习卡片翻转后接收用户评分，再据此更新复习间隔和下一次复习日期。测验模块需要支持选择题和填空题两种模式，并允许用户在 CET-4、CET-6 或混合词库范围内生成指定数量题目；系统在提交测验后应返回得分、错题列表和历史详情，以便用户进行错题回顾。",
    )
    add_paragraph(
        document,
        "统计模块则需要基于 review_log、user_word_progress 和 word_meta 等数据，汇总展示待复习数量、连续学习天数、已掌握词数、总学习词数、词性分布、热力图和周报文本。当前 DashboardController 已实现 overview 与 stats 两类接口，前端看板页则通过 RingChart、WeeklyReport 和 PressureAlert 等组件展示这些数据，满足了学习反馈可视化需求。",
    )

    add_heading2(document, "3.3 非功能需求分析")
    add_paragraph(
        document,
        "安全性方面，系统要求所有受保护接口必须在登录后访问，并通过角色控制管理员功能边界。当前后端通过 Spring Security 与 JWT 过滤器完成身份识别，前端通过路由守卫限制未登录访问和管理员越权访问，这满足了基本的鉴权需求。对于模型相关配置，系统进一步区分全局模型与私有模型，也体现出一定的数据隔离意识。",
    )
    add_paragraph(
        document,
        "稳定性方面，AI 内容生成存在异步执行、超时失败和返回不完整等天然风险，因此系统必须具备重试、兜底、缓存和状态标记机制。WordServiceImpl 中对 pending、partial、fallback 等多种状态进行了处理，并通过 Redis 缓存稳定结果，说明当前项目对生成链路的异常控制已经纳入非功能设计范畴。除此之外，页面加载还通过骨架屏、轮询和错误提示保证交互的连续性。",
    )
    add_paragraph(
        document,
        "可维护性方面，系统要求代码结构清晰、职责分层明确、测试入口可执行。当前仓库前端按视图、组件、API、状态和工具函数分层，后端按控制层、服务层、数据访问层、实体和配置分层，并已包含若干单元测试、MockMvc 测试和集成测试文件。部署脚本和 Nginx 配置也为后期上线和运维提供了基础支持。",
    )

    add_heading2(document, "3.4 关键业务流程分析")
    add_paragraph(
        document,
        "系统关键业务流程可归纳为以下四类。第一类是词条学习流程：用户从词库列表进入详情页，浏览基础释义和 AI 生成内容后将单词加入学习计划，系统创建 user_word_progress 记录，并在后续复习中持续更新状态。第二类是 AI 生成流程：页面发起详情请求后，后端读取缓存与数据库状态，必要时调用异步服务生成内容，并在生成完成后刷新详情展示。",
    )
    add_paragraph(
        document,
        "第三类是复习调度流程：系统根据 next_review_date 选择当日应复习词汇，用户在卡片背面完成评分后，ReviewService 调用 SM2Algorithm 更新 easiness、interval 和 repetition，再同步写入 review_log 和看板缓存。第四类是模拟测验流程：QuizService 从词表中随机抽取词汇生成选择题或填空题，提交后记录测验详情与错题信息，以支持历史回看和学习反馈。",
    )
    add_heading2(document, "3.5 本章小结")
    add_paragraph(
        document,
        "本章围绕当前项目的真实实现完成了需求分析，明确了系统目标、用户角色、功能需求、非功能需求和关键业务流程。与仅停留在开题阶段的目标描述相比，当前需求分析更加贴近实际代码结构和运行行为，为后续的总体设计和详细实现分析提供了明确边界。",
    )

    add_page_break(document)
    add_heading1(document, "第四章 系统总体设计")
    add_heading2(document, "4.1 系统总体架构设计")
    add_paragraph(
        document,
        "系统整体采用浏览器/服务器模式和前后端分离架构。前端使用 Vue 3 构建用户交互界面，后端基于 Spring Boot 提供 API 服务，MySQL 负责业务数据持久化，Redis 负责热点缓存与看板缓存，大模型接入层同时支持本地 Ollama 与云端兼容接口。按照职责可进一步划分为表示层、业务层、数据层和模型适配层四部分。",
    )
    add_paragraph(
        document,
        "表示层主要包括 Dashboard、WordList、WordDetail、Review、Quiz、Assistant、Profile 和管理端页面。业务层则以 AuthController、WordController、ReviewController、QuizController、AssistantController、DashboardController 和各类 Service 实现为核心。数据层围绕 user、word_meta、user_word_progress、review_log、quiz_session_record、word_import_batch 等表组织持久化数据。模型适配层承担模型调用、配置解析、提示构造、返回解析和使用量记录等职责。",
    )

    add_heading2(document, "4.2 功能模块设计")
    add_heading3(document, "4.2.1 用户与权限模块")
    add_paragraph(
        document,
        "该模块负责注册、登录、登出、资料查看和角色边界控制。系统登录后向前端返回 JWT 令牌，前端将其保存在本地并在请求拦截器中统一携带。管理员角色可访问模型管理、词库导入和用户管理等后台能力，普通用户则仅能访问学习相关能力。这样的设计既满足了基础账号体系需求，也为后续扩展学习偏好、每日目标和模型风格配置提供了承载空间。",
    )
    add_heading3(document, "4.2.2 词汇学习模块")
    add_paragraph(
        document,
        "词汇学习模块由词库浏览、词条详情、加入学习和内容生成组成。词库浏览负责分页和筛选，词条详情负责展示基础释义与 AI 内容，加入学习负责初始化进度记录，AI 内容生成则围绕例句、近义词、助记和解释等字段展开。该模块既是系统的核心入口，也是后续复习与测验功能的数据来源。",
    )
    add_heading3(document, "4.2.3 复习与测验模块")
    add_paragraph(
        document,
        "复习模块的职责是根据学习进度组织当日复习队列，并在用户打分后执行进度更新；测验模块的职责是随机生成题目、收集作答结果、给出得分反馈并保存测验详情。两者共同承担“巩固记忆”的任务，但交互模式不同：复习更偏向过程性训练，测验更偏向阶段性检测。系统通过 review_log 和 quiz_session_record 分别记录这两类行为，为统计和回顾提供依据。",
    )
    add_heading3(document, "4.2.4 学习助手与模型管理模块")
    add_paragraph(
        document,
        "学习助手模块用于承接用户对词汇学习的延伸提问。它能够读取单词上下文、最近对话历史和用户风格设置，构造连续对话提示，并在需要时执行自动续写，以缓解部分模型输出被截断的问题。模型管理模块则负责维护全局模型和私有模型配置，支持模型协议、基础地址、访问路径和 API Key 的存储与启停控制。二者配合后，使系统从“固定词条生成”扩展到“可交互的学习支持”。",
    )

    add_heading2(document, "4.3 数据库设计")
    add_paragraph(
        document,
        "根据当前数据库脚本和迁移文件，系统数据库由词表、用户表、进度表、日志表、导入批次表和模型相关表共同构成。其中 `cet4zx` 与 `cet6zx` 保存基础词汇，`word_meta` 保存不同风格下生成的扩展内容，`user_word_progress` 保存复习调度参数，`review_log` 保存复习行为日志，`quiz_session_record` 保存测验历史，`word_import_batch` 与 `word_import_batch_item` 记录词库导入和回滚过程。",
    )
    add_paragraph(
        document,
        "这种设计遵循了“基础词表与扩展元数据分离”的思路。基础词表相对稳定，适合按 id 顺序管理；扩展元数据受模型风格、生成状态和异步回填影响较大，因此单独拆分到 `word_meta` 表更有利于后续重试、更新和缓存控制。学习进度和行为日志也分别建表，避免将调度参数与历史记录耦合在同一结构中，从而增强查询与维护的清晰度。",
    )

    add_heading2(document, "4.4 接口设计")
    add_paragraph(
        document,
        "系统接口设计遵循统一的 REST 风格和统一的结果封装。认证模块提供 `/auth/register`、`/auth/login` 和 `/auth/logout`；词汇模块提供 `/word/list`、`/word/detail`、`/word/learn/add` 和若干 AI 生成接口；复习模块提供 `/review/today`、`/review/submit` 和 `/review/session/progress`；测验模块提供 `/quiz/generate`、`/quiz/submit` 和历史查询接口；学习助手模块提供 `/assistant/chat` 和会话状态同步接口。统一的接口风格减少了前后端联调成本，也便于在管理端和用户端之间复用基础请求能力。",
    )

    add_heading2(document, "4.5 安全、缓存与异常处理设计")
    add_paragraph(
        document,
        "在安全设计方面，系统通过 JWT 和 Spring Security 过滤器对接口进行保护，并通过角色判断限制管理员接口访问。异常处理方面，后端提供统一结果结构和全局异常处理逻辑，前端则在请求层统一处理业务错误提示。缓存设计方面，系统将单词详情和看板概览作为典型热点数据，通过 Redis 减少数据库重复查询；对于 AI 生成结果，只有当内容达到稳定状态时才进入缓存，从而避免页面轮询期间读到陈旧数据。",
    )
    add_heading2(document, "4.6 本章小结")
    add_paragraph(
        document,
        "本章从总体架构、功能模块、数据库设计、接口设计以及安全与缓存策略等方面对系统整体方案进行了说明。当前项目的总体设计能够较好地支撑词汇学习、AI 生成、复习调度、测验反馈和模型管理等多类功能，并在工程层面体现出一定的可扩展性和可维护性。",
    )

    add_page_break(document)
    add_heading1(document, "第五章 系统详细设计与实现")
    add_heading2(document, "5.1 用户认证与权限控制实现")
    add_paragraph(
        document,
        "用户认证模块以 AuthController 和 AuthService 为入口，支持用户注册、登录和退出登录。系统在登录成功后生成 JWT 令牌，前端将其作为后续请求的身份凭据。路由守卫在进入受保护页面前检查令牌和用户信息，并根据角色决定是否允许访问管理端页面。这样的实现既保证了学习数据的私有性，也为管理员侧词库维护和模型管理提供了明确边界。",
    )
    add_paragraph(
        document,
        "当前项目进一步引入了基于角色的模型权限控制，例如私有云模型的创建、编辑、删除和启停权限，以及全局云模型的维护权限。这些能力在普通词汇学习系统中并不常见，但在本项目中具有现实意义，因为模型配置直接影响内容生成质量、可用性和成本控制。通过把这些能力纳入 RBAC 体系，系统避免了模型配置接口被无序调用的问题。",
    )

    add_heading2(document, "5.2 词库导入与词条管理实现")
    add_paragraph(
        document,
        "词库导入模块主要由管理端接口完成。管理员上传 CSV 文件后，系统先进行预览解析，统计新增、更新、跳过和错误条数，再在确认导入时写入批次表和批次明细表。若导入后发现问题，还可基于批次号执行回滚操作，删除新增词条或恢复被更新前的旧数据。此种设计相比简单覆盖式导入更加安全，也更符合词汇库持续维护的工程需求。",
    )
    add_paragraph(
        document,
        "普通用户侧则通过 `/word/list` 和 `/word/detail` 接口访问词条数据。WordController 首先对词库类型和分页参数进行标准化处理，随后交由 WordServiceImpl 从 `cet4zx`、`cet6zx` 或其他兼容词表中读取数据，并结合用户学习状态和词性信息进行组装。这一过程中，基础词表与扩展元数据被统一映射为前端可直接消费的响应对象，体现了后端对业务对象的再组织能力。",
    )

    add_heading2(document, "5.3 单词详情页与 AI 生成链路实现")
    add_paragraph(
        document,
        "单词详情页是系统最能体现“大模型解释能力”的页面之一。前端 WordDetail 视图负责展示音标、词性、基础释义、学习状态和各类 AI 内容，同时提供加入学习、跳转学习助手和重试生成等操作。页面加载时会先请求 `/word/detail`，若返回状态显示内容尚未生成完整，则自动启动轮询机制，等待异步生成任务写回结果。这种“先返回可用结构，再增量补齐内容”的设计避免了长时间阻塞页面渲染。",
    )
    add_paragraph(
        document,
        "后端 WordServiceImpl 在处理详情请求时，会优先读取 Redis 缓存；若缓存不存在或内容不稳定，则从词表和 `word_meta` 表组合基础数据与扩展数据。系统根据用户的风格偏好、模型提供方和所选模型，决定是否需要重新触发异步生成；对 pending、partial、fallback 等状态也分别进行补救与修正。通过这种细粒度状态机设计，项目将原本一次性的文本生成行为转化为可重试、可回填、可感知的服务流程。",
    )
    add_paragraph(
        document,
        "在生成内容结构上，系统将例句、近义词、助记内容和扩展解释拆分为多个字段，并辅以语法提示和分项状态标记。这样既有利于前端组件分区域展示，也为后续缓存命中、单项重试和内容审核提供了更灵活的控制粒度。这种实现方式比将模型输出整体保存为一段长文本更适合学习类应用场景。",
    )

    add_heading2(document, "5.4 学习助手与本地/云端双模型适配实现")
    add_paragraph(
        document,
        "学习助手模块用于承接词条详情之外的延伸学习需求。AssistantController 在收到用户问题后，会读取当前词汇上下文、最近对话记录和用户风格设置，然后基于 PromptTemplate 拼接系统提示、用户资料和当前问题。系统会根据用户选择的模型提供方决定调用 OllamaClient 还是 CloudLlmClient；若云端模型不可用，则自动回退到本地模型以保证基本服务不中断。",
    )
    add_paragraph(
        document,
        "为解决部分模型生成答案被截断的问题，学习助手模块在代码中增加了自动续写逻辑。系统会检查当前输出是否包含结束标记、是否存在可疑结尾以及是否需要构造 continuation prompt 继续生成，从而提升最终回答的完整性。该实现体现了当前项目在“模型调用成功”之外，还进一步考虑了“回答可读性与完整性”的工程细节。",
    )
    add_paragraph(
        document,
        "另外，模型配置的复杂度也被纳入实现范围。管理员和用户均可维护云端模型配置，系统能够根据协议类型、模型键名、基础地址和 API Key 解析运行时配置，并通过使用量统计模块对模型调用情况进行汇总。这说明项目已经从单一实验性接入，扩展到了更接近真实平台产品的模型管理形态。",
    )

    add_heading2(document, "5.5 复习调度与学习看板实现")
    add_paragraph(
        document,
        "复习模块前端采用翻卡式交互。Review 页面进入后先调用 `review/today` 获取待复习卡片列表，用户翻转卡片后可对掌握程度进行打分，分值将直接决定 SM-2 算法更新结果。前端不仅保存当前题目索引和翻转状态，还记录每张卡片的作答耗时，以便在一轮复习结束后给出本次复习汇总信息。这种设计使复习过程更贴近实际学习行为。",
    )
    add_paragraph(
        document,
        "后端 ReviewServiceImpl 负责读取当日队列、提交评分并更新学习进度。在用户提交评分后，系统会基于当前 easiness、interval 和 repetition 调用 SM2Algorithm 计算新的复习间隔与下次复习日期，再更新 `user_word_progress` 表，同时将此次行为写入 `review_log`。之后，系统还会刷新 `daily_plan_cache` 和看板缓存，以确保待复习数、完成率和连续学习统计能够实时反映最新学习结果。",
    )
    add_paragraph(
        document,
        "看板模块则基于 `review_log`、`user_word_progress` 和 `word_meta` 构建概览数据。DashboardController 会统计待复习数量、连续学习天数、已掌握词数、总学习词数、压力指数以及近一段时间内的复习趋势和词性分布。前端通过图环、周报文本和压力提示等方式呈现这些数据，使用户能够从结果层面感知复习效果，而不仅仅停留在逐个背单词的过程视角。",
    )

    add_heading2(document, "5.6 模拟测验与历史记录实现")
    add_paragraph(
        document,
        "测验模块的设计目标是为用户提供阶段性检验能力。QuizServiceImpl 支持根据题目数量、题型模式和词库范围随机抽取单词，生成选择题或填空题，并将正确答案、单词信息和题目编号组织成会话对象保存在内存会话存储中。提交时，系统比对用户答案与正确答案，计算总题数、正确题数和错题列表，并将题目级详情序列化后写入 `quiz_session_record` 表。",
    )
    add_paragraph(
        document,
        "前端 Quiz 页面则围绕测验设置、题目作答、结果反馈和历史详情四个阶段组织交互。用户可以选择 10、20 或 30 题，指定 CET-4、CET-6 或混合词库，并在答题完成后查看错题和历史记录。相比单纯的“立即判分”，系统还支持点击历史记录查看题目级作答详情，从而为用户后续复盘和针对性训练提供依据。",
    )

    add_heading2(document, "5.7 本章小结")
    add_paragraph(
        document,
        "本章围绕当前仓库中的核心实现，对用户认证、词库导入、单词详情与 AI 生成、学习助手、复习调度、学习看板和模拟测验等关键模块进行了详细分析。与任务书和开题报告相比，当前系统在模型管理、状态控制、缓存回填和历史记录方面已经形成了较完整的工程实现，这也是本文后续测试与部署分析的重要基础。",
    )

    add_page_break(document)
    add_heading1(document, "第六章 系统测试与部署")
    add_heading2(document, "6.1 测试环境与测试策略")
    add_paragraph(
        document,
        "当前项目测试环境以本地开发环境为主，后端技术栈为 Spring Boot 3.3、JDK 17、Maven 和 MySQL/Redis，前端技术栈为 Vue 3、Vite 和 Element Plus。测试策略采用“自动化校验 + 手工用例整理”的组合方式：自动化部分覆盖算法逻辑、用户信息服务、角色权限、云模型边界场景和前端生产构建；手工部分则围绕模型配置、页面交互和主题切换等功能整理了测试用例。",
    )
    add_paragraph(
        document,
        "为了保证论文中的测试描述基于真实仓库证据，本文优先引用当前 `docs/test-cases-2026-04-05.md` 中已经记录的用例和验证结果，并补充对本地构建与测试命令的重新执行情况。这样的处理方式虽然不等同于大规模用户实测，但能够确保测试章节建立在当前项目快照的真实状态之上，而不是凭空构造的理想化结果。",
    )

    add_heading2(document, "6.2 自动化测试结果分析")
    add_paragraph(
        document,
        "在自动化测试方面，后端选取了 `SM2AlgorithmTest`、`UserControllerMockMvcTest`、`UserServiceImplTest`、`CloudLlmModelServiceImplBoundaryTest` 和 `RolePermissionServiceTest` 等用例进行回归执行。测试运行过程中，`UserServiceImplTest` 会主动模拟 `review_log` 查询失败场景并验证服务降级逻辑，因此控制台存在告警日志，但整体测试命令仍已通过，说明系统在部分依赖异常时具备兜底能力。",
    )
    add_paragraph(
        document,
        "其中，SM-2 算法测试覆盖了低分重置、前两次成功复习的固定间隔更新、E-Factor 下限约束和下一次复习日期计算等核心逻辑；云模型边界测试覆盖了私有模型更新、跨用户访问限制、模型键长度边界和清理 API Key 时的凭证删除逻辑。测试执行前，本文对 `CloudLlmModelServiceImplBoundaryTest` 中因 MyBatis-Plus 重载而出现的测试编译歧义进行了极小范围的测试代码兼容修正，使其能够在当前依赖版本下顺利执行。",
    )
    add_paragraph(
        document,
        "前端方面，通过执行 `npm run build` 完成了生产环境构建校验，构建过程顺利完成并生成 dist 产物。构建输出显示多个页面和组件资源被成功打包，说明当前登录、注册、看板、词库、详情、复习、测验、学习助手和管理页面至少在静态编译层面不存在阻断性错误。相较仅靠截图证明“页面存在”，生产构建通过更能说明当前前端工程具备可发布性。",
    )

    add_heading2(document, "6.3 功能测试与结果分析")
    add_paragraph(
        document,
        "结合现有测试文档，系统已对云模型新增、自检、公私有模型显示区分、协议可选输入以及登录/注册页主题切换等功能编写了较完整的手工测试用例。虽然受当前执行环境限制，论文撰写阶段未对全部浏览器联调场景逐条重放，但这些用例已经明确了前置条件、执行步骤和预期结果，可直接作为后续联调和答辩前回归检查的依据。",
    )
    add_paragraph(
        document,
        "从学习功能角度分析，系统的主要业务闭环包括：词库浏览后加入学习、详情页异步生成 AI 内容、复习队列读取与评分提交、看板统计刷新，以及测验生成与历史记录保存。当前仓库中对应的控制器、服务实现、数据表和前端页面均已具备，说明这些功能已经从设计稿阶段进入了可运行的工程状态。自动化测试和构建校验结果也进一步支持了这一判断。",
    )

    add_heading2(document, "6.4 部署过程与运行支撑")
    add_paragraph(
        document,
        "根据仓库中的 `deploy.sh` 脚本，系统部署流程包括拉取最新代码、前端构建、后端打包、停止旧进程、启动新 Jar 和重载 Nginx 六个步骤。前端构建结果位于 dist 目录，后端以可执行 Jar 形式运行，Nginx 负责对前端静态资源进行托管，并将 `/api/` 路径代理到 `127.0.0.1:8080` 的 Spring Boot 服务。该部署方式结构清晰，适合在中小规模服务器环境中快速落地。",
    )
    add_paragraph(
        document,
        "部署设计的优势在于前端与后端解耦明显，前端更新主要体现在静态资源重新构建和替换，后端更新则主要体现为 Jar 包替换与进程重启；当模型接入方式、缓存策略或管理页面变化时，可通过不同比例地更新前后端组件完成迭代。同时，Nginx 的 history 模式路由回退配置也为单页应用多页面访问提供了必要支持。",
    )

    add_heading2(document, "6.5 测试结论与问题分析")
    add_paragraph(
        document,
        "综合自动化测试、前端构建和当前代码审查结果，可以认为该系统已经实现了任务书中要求的主要能力，并在若干方面超出了开题阶段的预设范围。其主要优势体现在：一是词库管理、AI 生成、复习调度、测验和看板功能已经形成闭环；二是本地模型与云端模型并存，提高了生成策略的灵活性；三是异步生成、缓存控制、状态修正和模型权限管理等工程细节较为完善。",
    )
    add_paragraph(
        document,
        "但系统也仍存在需要继续优化的方面。例如，当前测试更多依赖本地自动化和手工用例整理，缺少更长周期的大样本学习数据评估；测验模块的题目生成仍以词表随机抽样为主，尚未结合用户错题画像和掌握程度进行个性化命题；学习助手虽然已经具备续写和回退机制，但在复杂对话场景下仍可进一步增强上下文压缩与答案可解释性。这些问题将作为后续优化重点。",
    )
    add_heading2(document, "6.6 本章小结")
    add_paragraph(
        document,
        "本章基于当前仓库中的测试脚本、测试文档、构建结果和部署文件，对系统的测试与部署情况进行了分析。结果表明，项目在核心功能闭环、自动化回归、前端构建和部署路径上已经具备较好的工程基础，但在更高层次的效果评估和更复杂场景优化方面仍有提升空间。",
    )

    add_page_break(document)
    add_heading1(document, "结  论")
    add_paragraph(
        document,
        "本文以“集成大模型解释能力的四六级高频词记忆系统”为研究对象，在任务书和开题报告基础上，以当前项目代码为最高优先级，对系统的需求分析、总体设计、详细实现、测试与部署过程进行了系统梳理。研究结果表明，当前项目已经实现了词库管理、智能解释生成、SM-2 复习调度、模拟测验、学习看板、学习助手和模型管理等关键能力，形成了较为完整的英语词汇学习闭环。",
    )
    add_paragraph(
        document,
        "从工程实现角度看，本项目的特点主要体现在三个方面。第一，将大模型能力嵌入到结构化学习系统中，而不是停留在单纯的聊天问答层面，使例句、近义词、助记和解释等结果可以被前端稳定消费。第二，在本地模型与云端模型之间建立了统一的适配路径，并通过缓存、异步状态、重试和回退机制提升了系统可用性。第三，将复习调度、测验训练和看板统计与模型能力结合起来，使学习行为能够被持续记录和反馈。",
    )
    add_paragraph(
        document,
        "当然，当前系统仍然存在进一步完善空间。例如，可继续补充面向真实学习者的长期效果评估，细化图表与工程图的论文呈现方式，完善更丰富的提示词模板与错题驱动测验策略，并进一步优化大模型调用成本、响应时延和上下文利用效率。后续若能结合更多真实学习数据与用户反馈，对复习策略和内容生成质量进行迭代，系统在教学辅助和个性化学习方面仍有较大的拓展前景。",
    )

    references = [
        "[1] 卢宇, 余京蕾, 陈鹏鹤, 等. 多模态大模型的教育应用研究与展望[J]. 电化教育研究, 2023, 44(6): 38-44.",
        "[2] 王艾艾. 语料库数据驱动下的英语四级词汇教学研究[J]. Advances in Education, 2024, 14: 1-10.",
        "[3] 张艺凡, 赵静怡, 藕才俊. 基于Ebbinghaus模型的高效记忆工具设计与实现[J]. 计算机科学与应用, 2024, 14(10): 22-32.",
        "[4] 陈义嘉, 刘芳. 基于大模型技术的个性化外语学习系统[J]. Software Engineering and Applications, 2025, 14: 647-658.",
        "[5] 方明炜, 朱君辉, 鲁鹿鸣, 等. 例句质量评估体系构建及大语言模型例句生成能力评估[J]. 2025.",
        "[6] 郭子浩, 孙由之, 张梦林, 等. 大语言模型背景下提示词工程赋能英语口语学习研究[J]. Advances in Education, 2023, 13: 8213-8220.",
        "[7] White J, Fu Q, Hays P, et al. A Prompt Pattern Catalog to Enhance Prompt Engineering with ChatGPT[J/OL]. arXiv, 2023.",
        "[8] Wang S, Hu Y, Yang X, et al. Personalized Forgetting Mechanism with Concept-Driven Knowledge Tracing[J]. 2024.",
        "[9] Subramani S, Anto J S, Sarensanth C, et al. Learn Buddy: Transforming Education with Generative AI[C]. 2024 International Conference on Computing and Data Science, 2024: 1-5.",
        "[10] Fincham N X, Alvarez A A. Using Large Language Models to Facilitate L2 Proficiency Development through Personalized Feedback and Scaffolding[C]. Proceedings of the International CALL Research Conference, 2024: 59-64.",
        "[11] Spring Team. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/current/reference/html/, 2026-04-06.",
        "[12] Vue Team. Vue Router Documentation[EB/OL]. https://router.vuejs.org/, 2026-04-06.",
    ]
    add_page_break(document)
    add_references(document, references)
    add_ack(document)


def main():
    document = Document(SOURCE_DOCX)

    # Keep cover and declaration pages, rewrite title-related placeholders.
    set_paragraph_text(document.paragraphs[6], TITLE, size_pt=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    set_paragraph_text(document.paragraphs[7], "", size_pt=12, first_line_indent=False)
    set_paragraph_text(document.paragraphs[16], "提交日期      年   月   日", size_pt=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

    # Remove old thesis body and appended format instructions.
    for paragraph in list(document.paragraphs[32:])[::-1]:
        delete_paragraph(paragraph)

    add_content(document)
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
